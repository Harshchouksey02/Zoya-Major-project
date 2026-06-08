"""
Major Project Report Generator
Student: Zoya Ibrahim (0210CA241134)
Project: AgroVeda – E-Commerce Platform for Agricultural Products
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

doc = Document()

# ============================================================
# GLOBAL STYLES & HELPERS
# ============================================================

def set_cell_shading(cell, color_hex):
    """Set cell background shading."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_border(cell, **kwargs):
    """Set cell border."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", "4")}" w:space="0" '
            f'w:color="{val.get("color", "000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)

def add_formatted_paragraph(doc_or_cell, text, font_name='Arial', font_size=11,
                             bold=False, italic=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                             space_before=0, space_after=6, line_spacing=1.15,
                             color=None, underline=False, first_line_indent=None):
    """Add a formatted paragraph."""
    p = doc_or_cell.add_paragraph()
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.underline = underline
    if color:
        run.font.color.rgb = RGBColor(*color)
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    if first_line_indent:
        p.paragraph_format.first_line_indent = Inches(first_line_indent)
    return p

def add_multi_run_paragraph(doc_obj, runs_data, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                             space_before=0, space_after=6, line_spacing=1.15):
    """Add a paragraph with multiple runs of different formatting."""
    p = doc_obj.add_paragraph()
    for rd in runs_data:
        run = p.add_run(rd.get('text', ''))
        run.font.name = rd.get('font_name', 'Arial')
        run.font.size = Pt(rd.get('font_size', 11))
        run.font.bold = rd.get('bold', False)
        run.font.italic = rd.get('italic', False)
        run.font.underline = rd.get('underline', False)
        if rd.get('color'):
            run.font.color.rgb = RGBColor(*rd['color'])
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), rd.get('font_name', 'Arial'))
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    return p

def add_chapter_title(doc_obj, number, title):
    """Add a chapter title with proper formatting."""
    add_formatted_paragraph(doc_obj, '', font_size=11, space_after=0)
    text = f"{number}. {title}" if number else title
    add_formatted_paragraph(doc_obj, text, font_name='Arial', font_size=18,
                           bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                           space_before=6, space_after=12, line_spacing=1.5)

def add_sub_heading(doc_obj, text):
    """Add sub-heading (1.1, 2.1, etc.)."""
    add_formatted_paragraph(doc_obj, text, font_name='Arial', font_size=14,
                           bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                           space_before=12, space_after=6, line_spacing=1.5)

def add_sub_sub_heading(doc_obj, text):
    """Add sub-sub-heading (1.1.1, etc.)."""
    add_formatted_paragraph(doc_obj, text, font_name='Arial', font_size=12,
                           bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                           space_before=8, space_after=4, line_spacing=1.5)

def add_body_text(doc_obj, text, space_after=6):
    """Add body text."""
    return add_formatted_paragraph(doc_obj, text, font_name='Arial', font_size=11,
                           alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                           space_before=0, space_after=space_after, line_spacing=1.15)

def add_bullet(doc_obj, text, indent_level=0):
    """Add bullet point."""
    p = add_formatted_paragraph(doc_obj, '', font_name='Arial', font_size=11,
                               alignment=WD_ALIGN_PARAGRAPH.LEFT,
                               space_before=0, space_after=3, line_spacing=1.15)
    p.clear()
    prefix = "    " * indent_level + "• "
    run = p.add_run(prefix + text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    return p

def add_bold_bullet(doc_obj, label, text, indent_level=0):
    """Add bullet point with bold label."""
    p = doc_obj.add_paragraph()
    prefix = "    " * indent_level + "• "
    run1 = p.add_run(prefix + label)
    run1.font.name = 'Arial'
    run1.font.size = Pt(11)
    run1.font.bold = True
    run1._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    run2 = p.add_run(text)
    run2.font.name = 'Arial'
    run2.font.size = Pt(11)
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    return p

def create_styled_table(doc_obj, headers, rows, col_widths=None):
    """Create a styled table with headers and rows."""
    table = doc_obj.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Style header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.font.bold = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, "D9E2F3")

    # Style data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.name = 'Arial'
            run.font.size = Pt(10)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Set column widths
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Inches(width)

    # Add borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>') 
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)
    return table

def add_page_break(doc_obj):
    """Add a page break."""
    doc_obj.add_page_break()

def add_diagram_image(doc_obj, image_path, width_inches=5.5):
    """Add a centered diagram image."""
    p = doc_obj.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run()
    run.add_picture(image_path, width=Inches(width_inches))
    return p


# ============================================================
# PAGE SETUP
# ============================================================
section = doc.sections[0]
section.page_width = Inches(8.27)   # A4
section.page_height = Inches(11.69)  # A4
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

# Remove header/footer from first section (cover page etc.)
section.different_first_page_header_footer = False
header = section.header
header.is_linked_to_previous = False
footer = section.footer
footer.is_linked_to_previous = False

# ============================================================
# 1. COVER PAGE
# ============================================================
# Add empty lines for spacing
for _ in range(2):
    add_formatted_paragraph(doc, '', font_size=11, space_after=0)

add_formatted_paragraph(doc, 'SHRI RAM INSTITUTE OF TECHNOLOGY, JABALPUR',
                        font_name='Times New Roman', font_size=16, bold=True,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4, line_spacing=1.5)

add_formatted_paragraph(doc, '(Affiliated to RGPV, Bhopal)',
                        font_name='Times New Roman', font_size=12,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4, line_spacing=1.5)

add_formatted_paragraph(doc, 'Department of Computer Applications (MCA)',
                        font_name='Times New Roman', font_size=13, bold=True,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12, line_spacing=1.5)

# Add empty line
add_formatted_paragraph(doc, '', font_size=11, space_after=6)

add_formatted_paragraph(doc, 'MAJOR PROJECT REPORT',
                        font_name='Cambria', font_size=18, bold=True,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6, line_spacing=1.5)

add_formatted_paragraph(doc, 'on',
                        font_name='Times New Roman', font_size=14,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6, line_spacing=1.5)

add_formatted_paragraph(doc, '"AgroVeda – E-Commerce Platform for Agricultural Products"',
                        font_name='Cambria', font_size=16, bold=True,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=18,
                        line_spacing=1.5, color=(0, 51, 102))

# Add empty lines
for _ in range(2):
    add_formatted_paragraph(doc, '', font_size=11, space_after=6)

add_formatted_paragraph(doc, 'Submitted in partial fulfillment of the requirements',
                        font_name='Times New Roman', font_size=12,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, line_spacing=1.5)

add_formatted_paragraph(doc, 'for the degree of Master of Computer Applications (MCA)',
                        font_name='Times New Roman', font_size=12,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12, line_spacing=1.5)

add_formatted_paragraph(doc, '', font_size=6, space_after=6)

add_formatted_paragraph(doc, 'Submitted by:',
                        font_name='Times New Roman', font_size=13, bold=True,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4, line_spacing=1.5)

add_formatted_paragraph(doc, 'ZOYA IBRAHIM',
                        font_name='Times New Roman', font_size=14, bold=True,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, line_spacing=1.5)

add_formatted_paragraph(doc, 'Enrollment No: 0210CA241134',
                        font_name='Times New Roman', font_size=12,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12, line_spacing=1.5)

add_formatted_paragraph(doc, '', font_size=6, space_after=6)

add_formatted_paragraph(doc, 'Under the Guidance of:',
                        font_name='Times New Roman', font_size=13, bold=True,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4, line_spacing=1.5)

add_formatted_paragraph(doc, 'Prof. Sachin Dubey',
                        font_name='Times New Roman', font_size=14, bold=True,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, line_spacing=1.5)

add_formatted_paragraph(doc, 'Department of MCA, SRIT, Jabalpur',
                        font_name='Times New Roman', font_size=12,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=18, line_spacing=1.5)

add_formatted_paragraph(doc, '', font_size=6, space_after=6)

add_formatted_paragraph(doc, '4th Semester MCA | Session 2025-26',
                        font_name='Times New Roman', font_size=12, bold=True,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6, line_spacing=1.5)

add_page_break(doc)


# ============================================================
# 2. CERTIFICATE (Principal)
# ============================================================
for _ in range(2):
    add_formatted_paragraph(doc, '', font_size=11, space_after=0)

add_formatted_paragraph(doc, 'CERTIFICATE',
                        font_name='Times New Roman', font_size=15, bold=True,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=18,
                        line_spacing=1.5, underline=True)

add_formatted_paragraph(doc, '', font_size=6, space_after=6)

add_formatted_paragraph(
    doc,
    'This is to certify that Ms. Zoya Ibrahim, Enrollment No. 0210CA241134, '
    'a student of 4th Semester Master of Computer Applications (MCA) at '
    'Shri Ram Institute of Technology, Jabalpur (affiliated to RGPV, Bhopal), '
    'has successfully completed her Major Project titled '
    '"AgroVeda – E-Commerce Platform for Agricultural Products" '
    'during the academic session 2025-26.',
    font_name='Times New Roman', font_size=14,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=12, line_spacing=1.5)

add_formatted_paragraph(
    doc,
    'The project work has been carried out under the guidance of Prof. Sachin Dubey, '
    'Department of MCA, SRIT, Jabalpur. The work presented in this report is genuine '
    'and has been completed as per the requirements of the university curriculum.',
    font_name='Times New Roman', font_size=14,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=12, line_spacing=1.5)

add_formatted_paragraph(
    doc,
    'I wish her all the best for her future endeavors.',
    font_name='Times New Roman', font_size=14,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=24, line_spacing=1.5)

# Signature area
for _ in range(4):
    add_formatted_paragraph(doc, '', font_size=11, space_after=6)

add_formatted_paragraph(doc, 'Prof. Nidhi Tiwari',
                        font_name='Times New Roman', font_size=14, bold=True,
                        alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=2, line_spacing=1.5)

add_formatted_paragraph(doc, 'Principal',
                        font_name='Times New Roman', font_size=14,
                        alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=2, line_spacing=1.5)

add_formatted_paragraph(doc, 'Shri Ram Institute of Technology, Jabalpur',
                        font_name='Times New Roman', font_size=14,
                        alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=6, line_spacing=1.5)

add_page_break(doc)


# ============================================================
# 3. CERTIFICATE (Examiner)
# ============================================================
for _ in range(2):
    add_formatted_paragraph(doc, '', font_size=11, space_after=0)

add_formatted_paragraph(doc, 'CERTIFICATE',
                        font_name='Times New Roman', font_size=15, bold=True,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=18,
                        line_spacing=1.5, underline=True)

add_formatted_paragraph(doc, '', font_size=6, space_after=6)

add_formatted_paragraph(
    doc,
    'This is to certify that Ms. Zoya Ibrahim, Enrollment No. 0210CA241134, '
    'a student of 4th Semester MCA at Shri Ram Institute of Technology, Jabalpur, '
    'has presented and demonstrated her Major Project titled '
    '"AgroVeda – E-Commerce Platform for Agricultural Products" '
    'for the examination. The project has been evaluated on the basis of its '
    'technical content, methodology, implementation, and presentation.',
    font_name='Times New Roman', font_size=14,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=18, line_spacing=1.5)

# Examiner table
exam_table = doc.add_table(rows=4, cols=3)
exam_table.alignment = WD_TABLE_ALIGNMENT.CENTER

exam_headers = ['', 'Name / Designation', 'Signature & Date']
exam_rows_data = [
    ['Internal Examiner', '', ''],
    ['External Examiner', '', ''],
    ['HOD (MCA Dept.)', '', ''],
]

for i, h in enumerate(exam_headers):
    cell = exam_table.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(h)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_shading(cell, "D9E2F3")

for r_idx, row_d in enumerate(exam_rows_data):
    for c_idx, val in enumerate(row_d):
        cell = exam_table.rows[r_idx + 1].cells[c_idx]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(val)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        if c_idx == 0:
            run.font.bold = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Set minimum height for signature space
        if c_idx == 2:
            cell.height = Inches(0.6)

# Add borders to examiner table
tbl = exam_table._tbl
tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>') 
borders = parse_xml(
    f'<w:tblBorders {nsdecls("w")}>'
    f'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
    f'<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
    f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
    f'<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
    f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
    f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
    f'</w:tblBorders>'
)
tblPr.append(borders)

add_page_break(doc)


# ============================================================
# 4. GUIDE CERTIFICATE
# ============================================================
for _ in range(2):
    add_formatted_paragraph(doc, '', font_size=11, space_after=0)

add_formatted_paragraph(doc, 'GUIDE CERTIFICATE',
                        font_name='Times New Roman', font_size=15, bold=True,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=18,
                        line_spacing=1.5, underline=True)

add_formatted_paragraph(doc, '', font_size=6, space_after=6)

add_formatted_paragraph(
    doc,
    'This is to certify that the Major Project titled '
    '"AgroVeda – E-Commerce Platform for Agricultural Products" '
    'submitted by Ms. Zoya Ibrahim, Enrollment No. 0210CA241134, '
    'in partial fulfillment of the requirements for the degree of '
    'Master of Computer Applications (MCA) from Rajiv Gandhi Proudyogiki Vishwavidyalaya (RGPV), '
    'Bhopal, has been completed under my direct supervision and guidance.',
    font_name='Times New Roman', font_size=14,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=12, line_spacing=1.5)

add_formatted_paragraph(
    doc,
    'The work presented in this project report is original and has not been submitted '
    'elsewhere for the award of any other degree or diploma. The student has shown '
    'sincere dedication, technical competence, and a methodical approach throughout '
    'the development of this project. The project demonstrates the practical application '
    'of modern web technologies including React, TypeScript, Node.js, and PostgreSQL '
    'to solve a real-world problem in the agricultural e-commerce domain.',
    font_name='Times New Roman', font_size=14,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=12, line_spacing=1.5)

add_formatted_paragraph(
    doc,
    'I recommend this project for evaluation and approval.',
    font_name='Times New Roman', font_size=14,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=24, line_spacing=1.5)

for _ in range(4):
    add_formatted_paragraph(doc, '', font_size=11, space_after=6)

add_formatted_paragraph(doc, 'Prof. Sachin Dubey',
                        font_name='Times New Roman', font_size=14, bold=True,
                        alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=2, line_spacing=1.5)

add_formatted_paragraph(doc, 'Project Guide',
                        font_name='Times New Roman', font_size=14,
                        alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=2, line_spacing=1.5)

add_formatted_paragraph(doc, 'Department of MCA',
                        font_name='Times New Roman', font_size=14,
                        alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=2, line_spacing=1.5)

add_formatted_paragraph(doc, 'Shri Ram Institute of Technology, Jabalpur',
                        font_name='Times New Roman', font_size=14,
                        alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=6, line_spacing=1.5)

add_page_break(doc)


# ============================================================
# 5. SELF CERTIFICATE
# ============================================================
for _ in range(2):
    add_formatted_paragraph(doc, '', font_size=11, space_after=0)

add_formatted_paragraph(doc, 'SELF CERTIFICATE',
                        font_name='Times New Roman', font_size=15, bold=True,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=18,
                        line_spacing=1.5, underline=True)

add_formatted_paragraph(doc, '', font_size=6, space_after=6)

add_formatted_paragraph(
    doc,
    'I, Zoya Ibrahim, Enrollment No. 0210CA241134, a student of 4th Semester '
    'Master of Computer Applications (MCA) at Shri Ram Institute of Technology, '
    'Jabalpur (affiliated to RGPV, Bhopal), hereby declare that the Major Project '
    'titled "AgroVeda – E-Commerce Platform for Agricultural Products" '
    'is my original work carried out during the academic session 2025-26.',
    font_name='Times New Roman', font_size=14,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=12, line_spacing=1.5)

add_formatted_paragraph(
    doc,
    'The project has been developed under the able guidance of Prof. Sachin Dubey, '
    'Department of MCA, SRIT, Jabalpur. I certify that the work presented in this report '
    'has not been copied from any source without proper citation and has not been '
    'submitted previously or concurrently to any other institution for the award of '
    'any degree, diploma, or certificate.',
    font_name='Times New Roman', font_size=14,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=12, line_spacing=1.5)

add_formatted_paragraph(
    doc,
    'I take full responsibility for the correctness and authenticity of the content '
    'presented in this report. All the external resources and references used have been '
    'duly acknowledged in the References section.',
    font_name='Times New Roman', font_size=14,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=24, line_spacing=1.5)

for _ in range(4):
    add_formatted_paragraph(doc, '', font_size=11, space_after=6)

add_formatted_paragraph(doc, 'Zoya Ibrahim',
                        font_name='Times New Roman', font_size=14, bold=True,
                        alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=2, line_spacing=1.5)

add_formatted_paragraph(doc, 'Enrollment No: 0210CA241134',
                        font_name='Times New Roman', font_size=14,
                        alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=2, line_spacing=1.5)

add_formatted_paragraph(doc, '4th Semester MCA, SRIT, Jabalpur',
                        font_name='Times New Roman', font_size=14,
                        alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=6, line_spacing=1.5)

add_page_break(doc)


# ============================================================
# 6. ACKNOWLEDGEMENT
# ============================================================
for _ in range(2):
    add_formatted_paragraph(doc, '', font_size=11, space_after=0)

add_formatted_paragraph(doc, 'ACKNOWLEDGEMENT',
                        font_name='Times New Roman', font_size=15, bold=True,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=18,
                        line_spacing=1.5, underline=True)

add_formatted_paragraph(doc, '', font_size=6, space_after=6)

add_formatted_paragraph(
    doc,
    'The completion of this Major Project would not have been possible without the '
    'guidance, support, and encouragement of several individuals. I take this opportunity '
    'to express my sincere gratitude to all those who have contributed to the successful '
    'completion of this project.',
    font_name='Times New Roman', font_size=14,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=12, line_spacing=1.5)

add_formatted_paragraph(
    doc,
    'First and foremost, I would like to express my deepest gratitude to my project guide, '
    'Prof. Sachin Dubey, Department of MCA, SRIT, Jabalpur, for his invaluable guidance, '
    'constant motivation, and expert supervision throughout the course of this project. '
    'His insightful suggestions and constructive feedback were instrumental in shaping '
    'this project to its present form.',
    font_name='Times New Roman', font_size=14,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=12, line_spacing=1.5)

add_formatted_paragraph(
    doc,
    'I am immensely grateful to Prof. Nidhi Tiwari, Principal, Shri Ram Institute of '
    'Technology, Jabalpur, for providing the necessary infrastructure and academic '
    'environment that facilitated the smooth execution of this project work.',
    font_name='Times New Roman', font_size=14,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=12, line_spacing=1.5)

add_formatted_paragraph(
    doc,
    'I would also like to extend my heartfelt thanks to all the faculty members of the '
    'Department of MCA for their valuable teachings and encouragement throughout my '
    'academic journey at SRIT.',
    font_name='Times New Roman', font_size=14,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=12, line_spacing=1.5)

add_formatted_paragraph(
    doc,
    'A special word of thanks goes to my friends – Harsh Chouksey, Aashika Jain, and '
    'Monika Haldkar – for their unwavering support, stimulating discussions, and for '
    'being a constant source of motivation during the challenging phases of this project. '
    'Their camaraderie and technical insights enriched my learning experience.',
    font_name='Times New Roman', font_size=14,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=12, line_spacing=1.5)

add_formatted_paragraph(
    doc,
    'Last but not least, I express my profound gratitude to my beloved family for their '
    'unconditional love, moral support, patience, and constant encouragement that has '
    'been the foundation of my academic pursuits. Without their sacrifices and blessings, '
    'this achievement would not have been possible.',
    font_name='Times New Roman', font_size=14,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=18, line_spacing=1.5)

for _ in range(2):
    add_formatted_paragraph(doc, '', font_size=11, space_after=6)

add_formatted_paragraph(doc, 'Zoya Ibrahim',
                        font_name='Times New Roman', font_size=14, bold=True,
                        alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=2, line_spacing=1.5)

add_page_break(doc)


# ============================================================
# 7. DECLARATION
# ============================================================
for _ in range(2):
    add_formatted_paragraph(doc, '', font_size=11, space_after=0)

add_formatted_paragraph(doc, 'DECLARATION',
                        font_name='Times New Roman', font_size=15, bold=True,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=18,
                        line_spacing=1.5, underline=True)

add_formatted_paragraph(doc, '', font_size=6, space_after=6)

add_formatted_paragraph(
    doc,
    'I, Zoya Ibrahim, Enrollment No. 0210CA241134, a student of 4th Semester '
    'Master of Computer Applications (MCA) at Shri Ram Institute of Technology, '
    'Jabalpur (affiliated to Rajiv Gandhi Proudyogiki Vishwavidyalaya, Bhopal), '
    'hereby declare that the Major Project Report titled '
    '"AgroVeda – E-Commerce Platform for Agricultural Products" '
    'submitted by me is a record of my own original work.',
    font_name='Times New Roman', font_size=14,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=12, line_spacing=1.5)

add_formatted_paragraph(
    doc,
    'I further declare that this project work has been carried out under the '
    'guidance of Prof. Sachin Dubey, Department of MCA, SRIT, Jabalpur, and has not been '
    'submitted in part or full to any other university or institution for the award of '
    'any degree, diploma, fellowship, or any other similar title or recognition.',
    font_name='Times New Roman', font_size=14,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=12, line_spacing=1.5)

add_formatted_paragraph(
    doc,
    'The information and data presented in this report are authentic and correct to '
    'the best of my knowledge and belief. I shall be solely responsible for the '
    'correctness of the facts and opinions expressed in this report.',
    font_name='Times New Roman', font_size=14,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=18, line_spacing=1.5)

for _ in range(2):
    add_formatted_paragraph(doc, '', font_size=11, space_after=6)

add_formatted_paragraph(doc, 'Date: 23 May 2026',
                        font_name='Times New Roman', font_size=14,
                        alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=2, line_spacing=1.5)

add_formatted_paragraph(doc, 'Place: SRIT, Jabalpur',
                        font_name='Times New Roman', font_size=14,
                        alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=18, line_spacing=1.5)

for _ in range(2):
    add_formatted_paragraph(doc, '', font_size=11, space_after=6)

add_formatted_paragraph(doc, 'Zoya Ibrahim',
                        font_name='Times New Roman', font_size=14, bold=True,
                        alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=2, line_spacing=1.5)

add_formatted_paragraph(doc, 'Enrollment No: 0210CA241134',
                        font_name='Times New Roman', font_size=14,
                        alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=6, line_spacing=1.5)

add_page_break(doc)


# ============================================================
# 8. INDEX PAGE — New section with page numbering starting here
# ============================================================
# Add a new section for page numbering
new_section = doc.add_section()
new_section.page_width = Inches(8.27)
new_section.page_height = Inches(11.69)
new_section.top_margin = Inches(1)
new_section.bottom_margin = Inches(1)
new_section.left_margin = Inches(1)
new_section.right_margin = Inches(1)

# Setup page numbering starting from 1
sectPr = new_section._sectPr
pgNumType = parse_xml(f'<w:pgNumType {nsdecls("w")} w:start="1"/>')
sectPr.append(pgNumType)

# Add footer with page number
footer = new_section.footer
footer.is_linked_to_previous = False
footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
# Add page number field
run = footer_para.add_run()
fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
run._element.append(fldChar1)
run2 = footer_para.add_run()
instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
run2._element.append(instrText)
run3 = footer_para.add_run()
fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
run3._element.append(fldChar2)

# INDEX title
for _ in range(1):
    add_formatted_paragraph(doc, '', font_size=11, space_after=0)

add_formatted_paragraph(doc, 'INDEX',
                        font_name='Cambria', font_size=15, bold=True,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=18,
                        line_spacing=1.5, underline=True)

add_formatted_paragraph(doc, '', font_size=6, space_after=6)

# Index table
index_data = [
    ['1', 'Company Introduction', '1-2'],
    ['2', 'About Project', '3-4'],
    ['3', 'Objective', '5-5'],
    ['4', 'SDLC', '6-7'],
    ['5', 'Software Requirement', '8-9'],
    ['6', 'Hardware Requirement', '10-11'],
    ['7', 'Use Case Diagram', '12-14'],
    ['8', 'ER Diagram', '15-16'],
    ['9', 'Data Flow Diagram', '16-18'],
    ['10', 'Database and Table', '18-22'],
    ['11', 'I/O and Interface Design', '22-27'],
    ['12', 'Testing', '28-31'],
    ['13', 'Future Scope', '31-32'],
    ['14', 'Conclusion', '32-34'],
    ['15', 'References', '34-35'],
]

idx_table = create_styled_table(doc, ['S.No.', 'Topic', 'Pg. No.'], index_data,
                                col_widths=[0.8, 4.0, 1.2])

add_page_break(doc)


# ============================================================
# 9. CHAPTER 1: COMPANY INTRODUCTION
# ============================================================
add_chapter_title(doc, '1', 'Company Introduction')

add_sub_heading(doc, '1.1 About DataSans')

add_body_text(doc, 'DataSans is a forward-thinking technology company that specializes in delivering cutting-edge digital solutions across diverse industry verticals. Founded with the vision of leveraging data-driven innovation to solve complex business challenges, DataSans has established itself as a prominent player in the IT solutions and consulting space. The company operates at the intersection of software engineering, data analytics, and digital transformation, offering end-to-end services that empower organizations to achieve operational excellence and competitive advantage.')

add_body_text(doc, 'Headquartered in India, DataSans has built a reputation for fostering a culture of innovation and technical excellence. The company provides internship and training opportunities to aspiring technology professionals, enabling them to gain hands-on experience with real-world projects while contributing to the company\'s mission of driving digital progress. DataSans believes in nurturing talent through practical exposure to modern technologies and agile development practices.')

add_body_text(doc, 'The company\'s portfolio encompasses a wide range of services including custom software development, web and mobile application development, e-commerce solutions, cloud computing services, data analytics and business intelligence, machine learning and AI solutions, and IT consulting. DataSans serves clients ranging from startups to established enterprises, providing tailored solutions that address specific business needs and market demands.')

add_sub_heading(doc, '1.2 Company Vision')

add_body_text(doc, 'DataSans envisions becoming a global leader in technology innovation, creating impactful digital solutions that transform businesses and improve lives. The company is committed to staying at the forefront of technological advancement while maintaining the highest standards of quality, reliability, and customer satisfaction. Their vision encompasses building a collaborative ecosystem where technology professionals, businesses, and end-users benefit from seamless digital experiences.')

add_sub_heading(doc, '1.3 Company Mission')

add_body_text(doc, 'The mission of DataSans is to empower businesses with innovative technology solutions that drive growth, efficiency, and digital transformation. The company aims to bridge the gap between technology and business objectives by delivering scalable, secure, and user-centric applications. DataSans is dedicated to continuous learning, knowledge sharing, and fostering an environment where creativity and technical expertise converge to produce exceptional results.')

add_sub_heading(doc, '1.4 Services Offered')

add_body_text(doc, 'DataSans provides a comprehensive suite of technology services designed to meet the evolving needs of modern businesses:')

add_bullet(doc, 'Custom Software Development – Designing and building bespoke software applications tailored to specific business requirements using modern technology stacks.')
add_bullet(doc, 'Web Application Development – Creating responsive, scalable, and feature-rich web applications using frameworks such as React, Angular, Vue.js, Node.js, and Django.')
add_bullet(doc, 'Mobile Application Development – Developing cross-platform and native mobile applications for Android and iOS platforms.')
add_bullet(doc, 'E-Commerce Solutions – Building comprehensive online retail platforms with payment gateway integration, inventory management, and analytics.')
add_bullet(doc, 'Cloud Computing & DevOps – Implementing cloud infrastructure solutions on AWS, Azure, and GCP with CI/CD pipelines and containerized deployments.')
add_bullet(doc, 'Data Analytics & Business Intelligence – Providing data-driven insights through advanced analytics, dashboards, and reporting tools.')
add_bullet(doc, 'Machine Learning & AI – Developing intelligent systems that leverage machine learning algorithms and artificial intelligence to automate processes and enhance decision-making.')

add_sub_heading(doc, '1.5 Technology Expertise')

add_body_text(doc, 'DataSans maintains expertise across a broad spectrum of technologies and platforms. The company\'s technical team is proficient in:')

add_bullet(doc, 'Frontend Technologies: React.js, Angular, Vue.js, TypeScript, HTML5, CSS3, TailwindCSS, Bootstrap')
add_bullet(doc, 'Backend Technologies: Node.js, Express.js, Python (Django/Flask), Java (Spring Boot), .NET Core')
add_bullet(doc, 'Database Systems: PostgreSQL, MySQL, MongoDB, Redis, Firebase Firestore')
add_bullet(doc, 'Cloud Platforms: AWS (EC2, S3, Lambda), Microsoft Azure, Google Cloud Platform, Vercel, Netlify')
add_bullet(doc, 'DevOps Tools: Docker, Kubernetes, Jenkins, GitHub Actions, GitLab CI/CD')
add_bullet(doc, 'Version Control: Git, GitHub, GitLab, Bitbucket')

add_sub_heading(doc, '1.6 Company Culture and Values')

add_body_text(doc, 'DataSans fosters a culture of continuous learning, collaboration, and innovation. The company values transparency, accountability, and ethical business practices. Team members are encouraged to explore new technologies, participate in knowledge-sharing sessions, and contribute to open-source projects. The organizational structure promotes flat hierarchies and open communication, enabling team members to voice ideas and take ownership of their work.')

add_body_text(doc, 'The company also places a strong emphasis on work-life balance and employee well-being, recognizing that a motivated and healthy workforce is essential for sustained innovation and productivity. Regular team-building activities, hackathons, and tech talks are organized to keep the team engaged and inspired.')

add_sub_heading(doc, '1.7 Project Development Environment')

add_body_text(doc, 'During the development of the AgroVeda project, the following development environment and tools were utilized:')

add_bold_bullet(doc, 'Operating System: ', 'Windows 11 – Primary development environment providing a stable and feature-rich platform for full-stack web development.')
add_bold_bullet(doc, 'Code Editor: ', 'Visual Studio Code (VS Code) – A lightweight yet powerful source code editor with extensions for TypeScript, React, ESLint, Prettier, and integrated terminal support.')
add_bold_bullet(doc, 'Version Control: ', 'Git and GitHub – Distributed version control system used for source code management, branch management, and collaborative development with pull request workflows.')
add_bold_bullet(doc, 'Database: ', 'PostgreSQL (via Supabase/Neon) – A robust, open-source relational database management system used for production data storage, complemented by an In-Memory JSON fallback for development and testing environments.')
add_bold_bullet(doc, 'Language & Framework: ', 'TypeScript, React 18, Vite, Node.js, Express.js – A modern full-stack JavaScript/TypeScript technology stack providing type safety, fast development server, and efficient production builds.')
add_bold_bullet(doc, 'Deployment: ', 'Vercel – A cloud platform for static sites and serverless functions, used for deploying both the frontend React application and the serverless backend API endpoints.')

add_page_break(doc)


# ============================================================
# 10. CHAPTER 2: ABOUT PROJECT
# ============================================================
add_chapter_title(doc, '2', 'About Project')

add_sub_heading(doc, '2.1 Project Title')

add_body_text(doc, '"AgroVeda – E-Commerce Platform for Agricultural Products"')

add_sub_heading(doc, '2.2 Project Overview')

add_body_text(doc, 'AgroVeda is a comprehensive full-stack e-commerce web application specifically designed and developed for the distribution and sale of agricultural products. The platform serves as a digital marketplace that bridges the gap between agricultural product manufacturers, distributors, and end consumers – primarily farmers and retailers operating in rural and semi-urban India. The application facilitates seamless online browsing, filtering, and purchasing of a diverse range of agricultural products including fertilizers, pesticides, growth boosters, herbicides, and organic farming solutions.')

add_body_text(doc, 'The platform has been engineered using a modern technology stack comprising React 18 with TypeScript for the frontend, Node.js with Express.js for the backend API, and PostgreSQL for persistent data storage. The application incorporates several enterprise-grade features including secure user authentication using JSON Web Tokens (JWT) with bcrypt password hashing, a comprehensive product catalog with bilingual support in Hindi and English, a fully functional shopping cart system, order management with Razorpay payment gateway integration, and a secure administrative dashboard for inventory and order management.')

add_body_text(doc, 'AgroVeda follows a responsive, mobile-first design philosophy built with TailwindCSS and Shadcn/UI component library, ensuring an optimal user experience across devices ranging from smartphones to desktop computers. The application is deployed on Vercel, leveraging serverless architecture for scalability and cost-effectiveness.')

add_sub_heading(doc, '2.3 Problem Statement')

add_body_text(doc, 'Agricultural product distribution in rural India continues to face significant inefficiencies and challenges that hinder both farmers and retailers from accessing quality products at competitive prices. The existing distribution model relies predominantly on local dealers and brick-and-mortar stores, which presents multiple limitations:')

add_bullet(doc, 'Limited Product Variety – Local dealers typically stock a narrow range of products from select brands, restricting farmers\' access to newer, more effective agricultural solutions available in the market.')
add_bullet(doc, 'Lack of Price Transparency – Without a centralized platform for price comparison, farmers often pay inflated prices for agricultural inputs, as they have no mechanism to verify the competitiveness of pricing offered by their local dealers.')
add_bullet(doc, 'No Digital Ordering Capability – The absence of an online ordering system means that farmers must physically visit stores, often traveling significant distances in rural areas, leading to loss of productive time and increased transportation costs.')
add_bullet(doc, 'No Online Payment Options – The agricultural retail ecosystem in rural India largely operates on a cash-only basis, limiting the convenience and traceability of transactions.')
add_bullet(doc, 'Language Barriers – Most existing e-commerce platforms operate exclusively in English, creating accessibility challenges for Hindi-speaking farmers and retailers in central India.')
add_bullet(doc, 'Lack of Product Information – Farmers often lack detailed information about product specifications, usage instructions, and suitability for specific crops, leading to suboptimal purchase decisions.')

add_sub_heading(doc, '2.4 Proposed Solution')

add_body_text(doc, 'AgroVeda addresses the identified challenges by providing a responsive, user-friendly e-commerce platform specifically tailored for the agricultural products domain. The proposed solution encompasses:')

add_bullet(doc, 'A bilingual (Hindi/English) web platform accessible on any device with an internet connection, eliminating geographical constraints and reducing the need for physical store visits.')
add_bullet(doc, 'A comprehensive product catalog with detailed descriptions, usage instructions, ratings, and reviews in both Hindi and English, empowering farmers to make informed purchase decisions.')
add_bullet(doc, 'An integrated shopping cart system with quantity management, bulk offers display, and real-time pricing calculations.')
add_bullet(doc, 'Multiple payment options including Razorpay payment gateway for online payments (UPI, debit/credit cards, net banking) and Cash on Delivery (COD) for areas with limited digital payment adoption.')
add_bullet(doc, 'Order tracking functionality allowing customers to monitor the status of their orders from placement to delivery.')
add_bullet(doc, 'A secure admin panel enabling administrators to manage products, monitor orders, update inventory, and manage user roles efficiently.')

add_sub_heading(doc, '2.5 Key Features – Customer Side')

add_body_text(doc, 'The customer-facing interface of AgroVeda provides a rich set of features designed to enhance the shopping experience:')

add_bold_bullet(doc, 'Product Browsing with Category Filters: ', 'Customers can browse through the complete product catalog organized by categories such as fertilizers, pesticides, growth boosters, herbicides, and organic products. Advanced filtering capabilities allow users to narrow results by category, price range, and rating.')
add_bold_bullet(doc, 'Hindi/English Bilingual Support: ', 'All product names, descriptions, and usage instructions are available in both Hindi and English, catering to the linguistic preferences of the target user base in central India.')
add_bold_bullet(doc, 'Shopping Cart with Quantity Management: ', 'A persistent shopping cart system allows customers to add products, adjust quantities, remove items, and view real-time subtotal calculations before proceeding to checkout.')
add_bold_bullet(doc, 'Online Checkout with Razorpay/COD: ', 'The checkout process supports both online payment through Razorpay (supporting UPI, cards, and net banking) and Cash on Delivery for maximum flexibility.')
add_bold_bullet(doc, 'Order History Tracking: ', 'Registered customers can view their complete order history, including order status, payment details, and itemized receipts.')
add_bold_bullet(doc, 'Responsive Mobile-First Design: ', 'The interface is optimized for mobile devices using TailwindCSS responsive utilities and Shadcn/UI components, ensuring a seamless experience across screen sizes.')

add_sub_heading(doc, '2.6 Key Features – Admin Side')

add_body_text(doc, 'The administrative interface provides comprehensive management capabilities:')

add_bold_bullet(doc, 'Secure Admin Authentication: ', 'Role-based access control with JWT authentication ensures that only authorized administrators can access the admin panel. Admin credentials are stored with bcrypt password hashing.')
add_bold_bullet(doc, 'Product CRUD Management: ', 'Administrators can create, read, update, and delete products through an intuitive interface. Product attributes include name (English/Hindi), category, price, unit, usage instructions, description, image URL, and bulk offer information.')
add_bold_bullet(doc, 'Order Monitoring Dashboard: ', 'A centralized dashboard displays all orders with their current status, enabling administrators to update order status (pending, processing, shipped, delivered, cancelled) and monitor revenue metrics.')
add_bold_bullet(doc, 'User Role Management: ', 'Administrators can manage user accounts and assign roles (customer/admin), ensuring proper access control across the platform.')

add_sub_heading(doc, '2.7 Technology Stack')

add_body_text(doc, 'The AgroVeda platform is built on a modern, full-stack technology architecture:')

add_bold_bullet(doc, 'Frontend: ', 'React 18 + TypeScript + Vite – Providing a component-based UI architecture with static type checking and fast Hot Module Replacement (HMR) during development.')
add_bold_bullet(doc, 'Backend: ', 'Node.js + Express.js – A lightweight and efficient server framework handling RESTful API endpoints, authentication middleware, and payment processing.')
add_bold_bullet(doc, 'Database: ', 'PostgreSQL / In-Memory JSON – PostgreSQL serves as the production database for persistent data storage, while an in-memory JSON store provides a zero-configuration fallback for development and testing.')
add_bold_bullet(doc, 'Payments: ', 'Razorpay – India\'s leading payment gateway, integrated for secure online payment processing with support for UPI, debit/credit cards, and net banking.')
add_bold_bullet(doc, 'Deployment: ', 'Vercel – A cloud platform optimized for frontend frameworks and serverless functions, providing automatic deployments, SSL certificates, and global CDN distribution.')

add_sub_heading(doc, '2.8 Project Scope')

add_body_text(doc, 'AgroVeda is currently designed as a single-vendor agricultural e-commerce platform, where a single business entity manages the product catalog and fulfillment operations. The platform\'s architecture, however, has been designed with scalability in mind, allowing for future expansion into a multi-vendor marketplace model where multiple agricultural product suppliers can register, list their products, and manage their own inventories. The current scope includes user registration and authentication, product catalog management, shopping cart and checkout, payment processing, order management, and administrative functions.')

add_page_break(doc)


# ============================================================
# 11. CHAPTER 3: OBJECTIVE
# ============================================================
add_chapter_title(doc, '3', 'Objective')

add_sub_heading(doc, '3.1 Primary Objective')

add_body_text(doc, 'The primary objective of the AgroVeda project is to design, develop, and deploy a responsive, full-stack agricultural e-commerce platform using modern web technologies – specifically React.js for the frontend and Node.js with Express.js for the backend. The platform aims to digitize the agricultural product distribution process, providing farmers and retailers with convenient online access to a wide range of agricultural inputs while supporting bilingual interaction in Hindi and English.')

add_sub_heading(doc, '3.2 Functional Objectives')

add_body_text(doc, 'The following functional objectives were defined for the successful completion of the AgroVeda project:')

add_bold_bullet(doc, 'Product Catalog with Bilingual Support: ', 'To develop a comprehensive product catalog system that stores and displays agricultural products with detailed information including names, descriptions, usage instructions, pricing, ratings, and images in both Hindi and English languages.')
add_bold_bullet(doc, 'Cart and Checkout System: ', 'To implement a fully functional shopping cart system that allows users to add products, modify quantities, calculate totals with discount support, and proceed through a streamlined checkout process with delivery information capture.')
add_bold_bullet(doc, 'Razorpay Payment Integration: ', 'To integrate Razorpay payment gateway for secure online payment processing, supporting multiple payment methods including UPI, debit/credit cards, and net banking, with proper order verification and payment confirmation flows.')
add_bold_bullet(doc, 'Admin Dashboard for CRUD Operations: ', 'To build a secure, role-protected admin dashboard that enables administrators to perform Create, Read, Update, and Delete operations on products, manage orders, update order statuses, and oversee user accounts.')

add_sub_heading(doc, '3.3 Non-Functional Objectives')

add_body_text(doc, 'In addition to the functional goals, the following non-functional objectives were established to ensure the overall quality and reliability of the platform:')

add_bold_bullet(doc, 'Security (JWT, bcrypt, CORS): ', 'To implement robust security measures including JWT-based stateless authentication, bcrypt password hashing with salt rounds, Cross-Origin Resource Sharing (CORS) configuration, and input validation to protect against common web vulnerabilities such as SQL injection and XSS attacks.')
add_bold_bullet(doc, 'Usability (Responsive Design): ', 'To create an intuitive, accessible, and responsive user interface following mobile-first design principles, ensuring the application provides an optimal experience across devices including smartphones, tablets, and desktop computers.')
add_bold_bullet(doc, 'Performance (Vite Build Optimization): ', 'To leverage Vite\'s modern build tooling for optimized production bundles with code splitting, tree shaking, and asset optimization, ensuring fast initial page loads and smooth runtime performance.')

add_sub_heading(doc, '3.4 Expected Outcomes')

add_body_text(doc, 'Upon successful completion of the project, the following outcomes were expected:')

add_bullet(doc, 'A fully functional e-commerce web application deployed on Vercel with a custom domain, accessible to users worldwide.')
add_bullet(doc, 'A PostgreSQL database schema implementing proper relational design with tables for users, roles, products, cart items, orders, and order items.')
add_bullet(doc, 'A working Razorpay payment gateway integration enabling customers to complete online transactions securely.')
add_bullet(doc, 'A responsive, bilingual user interface providing seamless product browsing, cart management, and order placement.')
add_bullet(doc, 'A secure admin panel for product and order management with role-based access control.')
add_bullet(doc, 'Comprehensive documentation covering system architecture, database design, testing results, and deployment procedures.')

add_page_break(doc)


# ============================================================
# 12. CHAPTER 4: SDLC
# ============================================================
add_chapter_title(doc, '4', 'Software Development Life Cycle (SDLC)')

add_sub_heading(doc, '4.1 Methodology: Agile Development')

add_body_text(doc, 'The AgroVeda project was developed following the Agile Software Development methodology, which emphasizes iterative development, continuous feedback, and adaptive planning. Agile was chosen as the development methodology for several reasons:')

add_bullet(doc, 'The project requirements evolved during development as new insights were gained about user needs and technical possibilities.')
add_bullet(doc, 'The iterative nature of Agile allowed for rapid prototyping and incremental feature delivery, ensuring early visibility of progress.')
add_bullet(doc, 'Short development cycles (sprints) enabled quick identification and resolution of issues, reducing the risk of late-stage project failures.')
add_bullet(doc, 'Agile\'s emphasis on working software over comprehensive documentation aligned well with the project\'s goal of delivering a functional product within the academic timeline.')

add_body_text(doc, 'The development was organized into 6 sprints, each lasting 3-4 days, with a total project development timeline of approximately 20 working days. Each sprint culminated in a demonstrable increment of functionality that was reviewed and validated before proceeding to the next sprint.')

add_sub_heading(doc, '4.2 Sprint 1: Project Setup and Foundation (3 Days)')

add_body_text(doc, 'The first sprint focused on establishing the project foundation and development infrastructure:')

add_bold_bullet(doc, 'Frontend Scaffolding: ', 'Initialized the React application using Vite with TypeScript template (npm create vite@latest). Configured TypeScript compiler options, ESLint rules, and Prettier formatting. Set up TailwindCSS with custom theme configuration and integrated Shadcn/UI component library.')
add_bold_bullet(doc, 'Backend API Setup: ', 'Created the Express.js server with modular route structure. Configured CORS middleware, JSON body parser, and error handling middleware. Set up environment variable management using dotenv.')
add_bold_bullet(doc, 'Database Schema Design: ', 'Designed the relational database schema with six tables: users, user_roles, products, cart_items, orders, and order_items. Created SQL migration scripts for PostgreSQL and implemented an in-memory JSON fallback storage system for development environments.')
add_bold_bullet(doc, 'Project Structure: ', 'Organized the codebase into a clean directory structure with separate folders for components, pages, hooks, services, types, and API routes.')

add_sub_heading(doc, '4.3 Sprint 2: Homepage and Product Catalog (4 Days)')

add_body_text(doc, 'The second sprint delivered the core user-facing product browsing experience:')

add_bold_bullet(doc, 'Banner Slider: ', 'Implemented an auto-rotating promotional banner carousel on the homepage showcasing seasonal offers and featured product categories. The slider supports manual navigation with previous/next controls and dot indicators.')
add_bold_bullet(doc, 'Product Grid: ', 'Developed a responsive product grid component that displays product cards with images, names (Hindi/English), prices, ratings, and quick add-to-cart buttons. The grid adapts from single column on mobile to four columns on desktop.')
add_bold_bullet(doc, 'Category Filters: ', 'Built a sidebar filter panel with category checkboxes, price range slider, and rating filter. Filters are applied client-side for instant results with URL query parameter synchronization.')
add_bold_bullet(doc, 'Bilingual UI: ', 'Implemented a language toggle system allowing users to switch between Hindi and English display for product names, descriptions, and UI labels. Language preference is persisted in localStorage.')

add_sub_heading(doc, '4.4 Sprint 3: Authentication System (3 Days)')

add_body_text(doc, 'The third sprint implemented the user authentication and authorization system:')

add_bold_bullet(doc, 'JWT Login/Signup: ', 'Created user registration and login API endpoints with proper validation. Implemented JWT token generation with configurable expiration, token verification middleware, and automatic token refresh mechanisms.')
add_bold_bullet(doc, 'bcrypt Password Hashing: ', 'Integrated bcryptjs for secure password hashing with 10 salt rounds, ensuring that user passwords are never stored in plaintext. Password comparison is performed using bcrypt\'s built-in timing-safe comparison function.')
add_bold_bullet(doc, 'Role-Based Access Control: ', 'Implemented a user_roles table and middleware that checks user roles (customer/admin) before granting access to protected API endpoints. Admin routes are secured with additional authorization checks.')

add_sub_heading(doc, '4.5 Sprint 4: Cart and Checkout (4 Days)')

add_body_text(doc, 'The fourth sprint focused on the shopping cart and order placement functionality:')

add_bold_bullet(doc, 'Add to Cart: ', 'Developed cart management API endpoints (add, update quantity, remove, clear) and a persistent cart storage system linked to authenticated users. Guest users can browse but must login to add items to cart.')
add_bold_bullet(doc, 'Quantity Management: ', 'Implemented increment/decrement controls for each cart item with real-time subtotal recalculation. Stock availability checks prevent over-ordering.')
add_bold_bullet(doc, 'Order Placement: ', 'Built the checkout flow including customer information capture (name, address, phone), order summary review, discount code application, and order confirmation with unique order number generation.')

add_sub_heading(doc, '4.6 Sprint 5: Payment Integration (3 Days)')

add_body_text(doc, 'The fifth sprint implemented the payment processing capabilities:')

add_bold_bullet(doc, 'Razorpay Order Creation: ', 'Integrated Razorpay Node.js SDK for server-side order creation. When a customer initiates payment, the backend creates a Razorpay order with the calculated amount and returns the order ID to the frontend.')
add_bold_bullet(doc, 'Payment Verification: ', 'Implemented payment verification using Razorpay\'s signature verification mechanism. After successful payment, the backend verifies the payment signature using HMAC SHA256 to ensure payment authenticity before updating the order status.')
add_bold_bullet(doc, 'COD Fallback: ', 'Added Cash on Delivery as an alternative payment method for customers who prefer offline payment. COD orders are marked with a "pending_payment" status until delivery confirmation.')

add_sub_heading(doc, '4.7 Sprint 6: Admin Panel and Deployment (3 Days)')

add_body_text(doc, 'The final sprint completed the admin functionality and production deployment:')

add_bold_bullet(doc, 'Product Management: ', 'Built the admin product management interface with forms for creating new products and editing existing ones. The interface supports all product attributes including bilingual fields, pricing, categorization, and image URL management.')
add_bold_bullet(doc, 'Order Tracking: ', 'Developed an admin order management dashboard displaying all orders with filtering by status, date range, and search by order number. Admins can update order statuses and view detailed order information.')
add_bold_bullet(doc, 'User Role Management: ', 'Implemented user listing and role assignment functionality, allowing admins to promote customers to admin role or revoke admin privileges.')
add_bold_bullet(doc, 'Deployment to Vercel: ', 'Configured the project for Vercel deployment with proper build settings, environment variable configuration, and serverless function setup. Set up automatic deployments from the main GitHub branch.')

add_page_break(doc)


# ============================================================
# 13. CHAPTER 5: SOFTWARE REQUIREMENT
# ============================================================
add_chapter_title(doc, '5', 'Software Requirement')

add_sub_heading(doc, '5.1 Frontend Technologies')

add_body_text(doc, 'The following software technologies and libraries were employed for the frontend development of AgroVeda:')

sw_frontend = [
    ['HTML5', 'Markup Language', 'Semantic structure and content of web pages'],
    ['CSS3 / TailwindCSS', 'Styling Framework', 'Utility-first CSS framework for responsive design'],
    ['TypeScript / ES6+', 'Programming Language', 'Statically-typed superset of JavaScript'],
    ['React 18', 'UI Library', 'Component-based library for building user interfaces'],
    ['Vite 5.x', 'Build Tool', 'Fast development server and optimized production bundler'],
    ['Radix UI', 'Component Primitives', 'Unstyled, accessible UI component primitives'],
    ['Shadcn/UI', 'Component Library', 'Pre-built, customizable UI components built on Radix'],
    ['React Router v6', 'Routing Library', 'Client-side routing for single-page application navigation'],
    ['React Query (TanStack)', 'Data Fetching', 'Server state management with caching and synchronization'],
    ['Recharts', 'Charting Library', 'Composable charting library for admin dashboard analytics'],
    ['Lucide Icons', 'Icon Library', 'Open-source icon set with React components'],
]

create_styled_table(doc, ['Technology', 'Type', 'Purpose'], sw_frontend, col_widths=[1.8, 1.5, 3.0])

add_formatted_paragraph(doc, '', font_size=6, space_after=6)

add_sub_heading(doc, '5.2 Backend Technologies')

sw_backend = [
    ['Node.js 18+', 'Runtime Environment', 'JavaScript runtime for server-side execution'],
    ['Express.js 4.x', 'Web Framework', 'Minimalist web framework for building REST APIs'],
    ['bcryptjs', 'Security Library', 'Password hashing with salt for secure authentication'],
    ['jsonwebtoken', 'Auth Library', 'JWT generation and verification for stateless auth'],
    ['pg (node-postgres)', 'Database Client', 'PostgreSQL client for Node.js with query support'],
    ['cors', 'Middleware', 'Cross-Origin Resource Sharing configuration'],
    ['dotenv', 'Configuration', 'Environment variable management from .env files'],
]

create_styled_table(doc, ['Technology', 'Type', 'Purpose'], sw_backend, col_widths=[1.8, 1.5, 3.0])

add_formatted_paragraph(doc, '', font_size=6, space_after=6)

add_sub_heading(doc, '5.3 Database')

sw_db = [
    ['PostgreSQL 15+', 'Production Database', 'Relational database for persistent data storage'],
    ['In-Memory JSON', 'Dev Database', 'Zero-config fallback storage for development/testing'],
]

create_styled_table(doc, ['Technology', 'Type', 'Purpose'], sw_db, col_widths=[1.8, 1.5, 3.0])

add_formatted_paragraph(doc, '', font_size=6, space_after=6)

add_sub_heading(doc, '5.4 Payment Gateway')

add_body_text(doc, 'Razorpay SDK is integrated for processing online payments. Razorpay is India\'s leading payment gateway supporting UPI, debit/credit cards, net banking, and wallet payments. The integration includes server-side order creation using the Razorpay Node.js SDK and client-side checkout using Razorpay\'s JavaScript widget.')

add_sub_heading(doc, '5.5 Development Tools')

sw_tools = [
    ['Visual Studio Code', 'Code Editor', 'Primary IDE with TypeScript, React, and ESLint plugins'],
    ['Git', 'Version Control', 'Distributed version control for source code management'],
    ['GitHub', 'Repository Hosting', 'Remote repository, collaboration, and CI/CD integration'],
    ['npm', 'Package Manager', 'Node.js package manager for dependency management'],
    ['Postman', 'API Testing', 'REST API testing and documentation tool'],
]

create_styled_table(doc, ['Tool', 'Type', 'Purpose'], sw_tools, col_widths=[1.8, 1.5, 3.0])

add_formatted_paragraph(doc, '', font_size=6, space_after=6)

add_sub_heading(doc, '5.6 Operating System and Browser Requirements')

add_body_text(doc, 'Development Environment:')
add_bullet(doc, 'Windows 10/11 – Primary development operating system')
add_bullet(doc, 'Vercel – Production deployment platform (Linux-based serverless environment)')

add_body_text(doc, 'Browser Compatibility:')
add_bullet(doc, 'Google Chrome 100+ (Primary development and testing browser)')
add_bullet(doc, 'Mozilla Firefox 100+')
add_bullet(doc, 'Apple Safari 15+')
add_bullet(doc, 'Microsoft Edge 100+')

add_sub_heading(doc, '5.7 Build and Run Commands')

add_body_text(doc, 'The following commands are used to build and run the AgroVeda application:')

add_bold_bullet(doc, 'Development Server: ', 'npm run dev – Starts the Vite development server with HMR on port 5173')
add_bold_bullet(doc, 'Production Build: ', 'npm run build – Creates an optimized production build in the dist/ directory')
add_bold_bullet(doc, 'Backend API (Local): ', 'node api/local.js – Starts the Express.js API server on port 3001')
add_bold_bullet(doc, 'Database Migration: ', 'node api/migrate.js – Runs database migration scripts to set up PostgreSQL tables')

add_page_break(doc)


# ============================================================
# 14. CHAPTER 6: HARDWARE REQUIREMENT
# ============================================================
add_chapter_title(doc, '6', 'Hardware Requirement')

add_sub_heading(doc, '6.1 Development Hardware')

add_body_text(doc, 'The following hardware configuration was used during the development of the AgroVeda project:')

hw_dev = [
    ['Processor', 'Intel Core i5 / AMD Ryzen 5 (or equivalent)', 'Multi-core processor for efficient build compilation and development server operation'],
    ['RAM', '8 GB DDR4 (Minimum)', 'Sufficient memory for running code editor, development server, database, and browser simultaneously'],
    ['Storage', '256 GB SSD (Minimum)', 'SSD provides fast read/write speeds essential for npm package installation, Vite builds, and database operations'],
    ['Display', '14" FHD (1920x1080)', 'Full HD display for comfortable code editing and UI development with adequate screen real estate'],
    ['Network', 'Broadband Internet', 'Stable internet connection for npm package downloads, Git operations, and Vercel deployments'],
    ['Input Devices', 'Standard Keyboard & Mouse', 'Standard input devices for development work'],
]

create_styled_table(doc, ['Component', 'Specification', 'Justification'], hw_dev, col_widths=[1.3, 2.2, 3.0])

add_formatted_paragraph(doc, '', font_size=6, space_after=6)

add_sub_heading(doc, '6.2 Server Hardware (Production)')

add_body_text(doc, 'The AgroVeda application is deployed on Vercel\'s serverless infrastructure, which eliminates the need for dedicated physical server hardware. Vercel provides auto-scaling, global CDN distribution, and SSL certificates as part of its platform:')

hw_server = [
    ['Frontend Hosting', 'Vercel Edge Network', 'Static files served via global CDN with automatic SSL'],
    ['Backend API', 'Vercel Serverless Functions', 'Auto-scaling serverless compute for API endpoints'],
    ['Database', 'PostgreSQL Cloud Instance', '1 GB storage allocation with automated backups via Supabase/Neon'],
    ['CDN', 'Vercel Edge Network (Global)', 'Content delivery across 70+ global edge locations for low latency'],
]

create_styled_table(doc, ['Component', 'Platform', 'Details'], hw_server, col_widths=[1.5, 2.0, 3.0])

add_formatted_paragraph(doc, '', font_size=6, space_after=6)

add_sub_heading(doc, '6.3 Client Hardware Requirements')

add_body_text(doc, 'AgroVeda is a web-based application accessible through any modern web browser. The minimum client-side hardware requirements are:')

hw_client = [
    ['Device', 'Smartphone / Tablet / Laptop / Desktop', 'Any device with a modern web browser'],
    ['Processor', 'Any modern processor (2015 or later)', 'Basic processing power for web application rendering'],
    ['RAM', '2 GB (Minimum)', 'Sufficient for browser-based application operation'],
    ['Storage', 'No local storage required', 'Application runs entirely in the browser; no installation needed'],
    ['Display', 'Any resolution (320px minimum width)', 'Responsive design adapts to all screen sizes'],
    ['Network', 'Internet Connection (3G or better)', 'Required for accessing the web application and API'],
]

create_styled_table(doc, ['Component', 'Requirement', 'Notes'], hw_client, col_widths=[1.3, 2.5, 2.7])

add_page_break(doc)


# ============================================================
# 15. CHAPTER 7: USE CASE DIAGRAM
# ============================================================
add_chapter_title(doc, '7', 'Use Case Diagram')

add_sub_heading(doc, '7.1 Introduction to Use Case Modeling')

add_body_text(doc, 'Use Case Diagrams are a fundamental component of the Unified Modeling Language (UML) and serve as a powerful tool for capturing the functional requirements of a system from the user\'s perspective. In the context of AgroVeda, use case diagrams illustrate the interactions between the system\'s actors (Customer and Admin) and the various functionalities provided by the platform. These diagrams help in understanding the scope of the system, identifying the key functionalities, and documenting the expected behavior of the application.')

add_sub_heading(doc, '7.2 System Actors')

add_body_text(doc, 'The AgroVeda system identifies two primary actors who interact with the platform:')

add_bold_bullet(doc, 'Customer: ', 'A registered or guest user who visits the platform to browse agricultural products, add items to the shopping cart, place orders, make payments, and track order status. Customers can register, login, manage their profile, and view order history.')
add_bold_bullet(doc, 'Admin: ', 'An authorized administrator who manages the backend operations of the platform including product catalog management, order processing, user role management, and monitoring platform analytics through the admin dashboard.')

add_sub_heading(doc, '7.3 Customer Use Cases')

cust_uc = [
    ['UC-C1', 'Browse Products', 'Customer browses the product catalog with category filters, search, and sorting options. Products are displayed in a responsive grid with images, prices, and ratings.'],
    ['UC-C2', 'Register / Login', 'Customer creates a new account with email and password or logs into an existing account using JWT-based authentication.'],
    ['UC-C3', 'Add to Cart', 'Customer adds selected products to the shopping cart with specified quantities. Cart persists across sessions for authenticated users.'],
    ['UC-C4', 'Place Order & Pay', 'Customer proceeds to checkout, provides delivery information, selects payment method (Razorpay/COD), and completes the order.'],
    ['UC-C5', 'Track Orders', 'Customer views order history and tracks the status of current orders (pending, processing, shipped, delivered).'],
]

create_styled_table(doc, ['Use Case ID', 'Use Case Name', 'Description'], cust_uc, col_widths=[1.0, 1.5, 4.0])

add_formatted_paragraph(doc, '', font_size=6, space_after=6)

add_sub_heading(doc, '7.4 Admin Use Cases')

admin_uc = [
    ['UC-A1', 'Admin Login', 'Admin authenticates using dedicated admin credentials with role-based JWT verification.'],
    ['UC-A2', 'Manage Products', 'Admin performs CRUD operations on products – add new products, edit details, update pricing/inventory, or remove discontinued products.'],
    ['UC-A3', 'Manage Orders', 'Admin views all orders, filters by status/date, updates order status (pending → processing → shipped → delivered), and handles order cancellations.'],
    ['UC-A4', 'Manage Users', 'Admin views registered users, manages user roles (customer/admin), and can deactivate accounts if necessary.'],
    ['UC-A5', 'View Dashboard', 'Admin accesses the analytics dashboard showing key metrics: total orders, revenue, product count, and recent order activity.'],
]

create_styled_table(doc, ['Use Case ID', 'Use Case Name', 'Description'], admin_uc, col_widths=[1.0, 1.5, 4.0])

add_formatted_paragraph(doc, '', font_size=12, space_after=6)

add_sub_heading(doc, '7.5 Use Case Diagram')

add_diagram_image(doc, 'diagrams/use_case_diagram.png', width_inches=5.8)
add_body_text(doc, 'Figure 7.1: Use Case Diagram for AgroVeda E-Commerce Platform showing interactions between Customer and Admin actors with system functionalities.')

add_page_break(doc)


# ============================================================
# 16. CHAPTER 8: ER DIAGRAM
# ============================================================
add_chapter_title(doc, '8', 'ER Diagram')

add_sub_heading(doc, '8.1 Introduction to Entity-Relationship Modeling')

add_body_text(doc, 'The Entity-Relationship (ER) Diagram is a visual representation of the data model underlying the AgroVeda application. It illustrates the entities (tables), their attributes (columns), and the relationships between them. The ER diagram serves as a blueprint for the database schema design, ensuring data integrity, normalization, and efficient query patterns. The AgroVeda database follows relational database principles with proper primary keys, foreign keys, and referential integrity constraints.')

add_sub_heading(doc, '8.2 Entities and Attributes')

add_body_text(doc, 'The AgroVeda database consists of six core entities:')

add_sub_sub_heading(doc, '8.2.1 Users Entity')
add_body_text(doc, 'The Users entity stores the authentication credentials and profile information for all registered users:')
add_bullet(doc, 'id (SERIAL, PRIMARY KEY) – Unique identifier for each user')
add_bullet(doc, 'email (VARCHAR, UNIQUE, NOT NULL) – User\'s email address used for login')
add_bullet(doc, 'password_hash (VARCHAR, NOT NULL) – bcrypt-hashed password')
add_bullet(doc, 'created_at (TIMESTAMP, DEFAULT NOW()) – Account creation timestamp')

add_sub_sub_heading(doc, '8.2.2 User_Roles Entity')
add_body_text(doc, 'The User_Roles entity implements role-based access control:')
add_bullet(doc, 'id (SERIAL, PRIMARY KEY) – Unique identifier')
add_bullet(doc, 'user_id (INTEGER, FOREIGN KEY → users.id) – Reference to the user')
add_bullet(doc, 'role (VARCHAR, NOT NULL) – Role name (\'customer\' or \'admin\')')

add_sub_sub_heading(doc, '8.2.3 Products Entity')
add_body_text(doc, 'The Products entity stores the complete product catalog with bilingual support:')
add_bullet(doc, 'id (SERIAL, PRIMARY KEY) – Unique product identifier')
add_bullet(doc, 'name (VARCHAR, NOT NULL) – Product name in English')
add_bullet(doc, 'name_hindi (VARCHAR) – Product name in Hindi')
add_bullet(doc, 'category (VARCHAR, NOT NULL) – Product category (fertilizer, pesticide, etc.)')
add_bullet(doc, 'price (DECIMAL, NOT NULL) – Product price in INR')
add_bullet(doc, 'unit (VARCHAR) – Unit of measurement (kg, litre, packet, etc.)')
add_bullet(doc, 'usage (TEXT) – Usage instructions in English')
add_bullet(doc, 'description (TEXT) – Detailed product description in English')
add_bullet(doc, 'description_hindi (TEXT) – Detailed product description in Hindi')
add_bullet(doc, 'rating (DECIMAL) – Average product rating (0-5)')
add_bullet(doc, 'reviews (INTEGER) – Total number of customer reviews')
add_bullet(doc, 'image_url (VARCHAR) – URL to the product image')
add_bullet(doc, 'bulk_offers (TEXT) – Bulk purchase offer information')
add_bullet(doc, 'created_at (TIMESTAMP) – Product creation timestamp')

add_sub_sub_heading(doc, '8.2.4 Cart_Items Entity')
add_body_text(doc, 'The Cart_Items entity manages the shopping cart for authenticated users:')
add_bullet(doc, 'id (SERIAL, PRIMARY KEY) – Unique cart item identifier')
add_bullet(doc, 'user_id (INTEGER, FOREIGN KEY → users.id) – Reference to the cart owner')
add_bullet(doc, 'product_id (INTEGER, FOREIGN KEY → products.id) – Reference to the product')
add_bullet(doc, 'quantity (INTEGER, NOT NULL, DEFAULT 1) – Quantity of the product in cart')

add_sub_sub_heading(doc, '8.2.5 Orders Entity')
add_body_text(doc, 'The Orders entity stores order information including customer details and payment status:')
add_bullet(doc, 'id (SERIAL, PRIMARY KEY) – Unique order identifier')
add_bullet(doc, 'user_id (INTEGER, FOREIGN KEY → users.id) – Reference to the ordering user')
add_bullet(doc, 'order_number (VARCHAR, UNIQUE) – Human-readable order number')
add_bullet(doc, 'customer_name (VARCHAR) – Delivery recipient name')
add_bullet(doc, 'customer_address (TEXT) – Delivery address')
add_bullet(doc, 'customer_phone (VARCHAR) – Contact phone number')
add_bullet(doc, 'subtotal (DECIMAL) – Order subtotal before discounts')
add_bullet(doc, 'discount_amount (DECIMAL) – Applied discount amount')
add_bullet(doc, 'total_amount (DECIMAL) – Final order total')
add_bullet(doc, 'payment_method (VARCHAR) – Payment method (razorpay/cod)')
add_bullet(doc, 'status (VARCHAR) – Order status (pending, processing, shipped, delivered, cancelled)')
add_bullet(doc, 'razorpay_order_id (VARCHAR) – Razorpay order reference ID')
add_bullet(doc, 'razorpay_payment_id (VARCHAR) – Razorpay payment reference ID')
add_bullet(doc, 'created_at (TIMESTAMP) – Order placement timestamp')

add_sub_sub_heading(doc, '8.2.6 Order_Items Entity')
add_body_text(doc, 'The Order_Items entity stores individual items within each order:')
add_bullet(doc, 'id (SERIAL, PRIMARY KEY) – Unique order item identifier')
add_bullet(doc, 'order_id (INTEGER, FOREIGN KEY → orders.id) – Reference to the parent order')
add_bullet(doc, 'product_id (INTEGER, FOREIGN KEY → products.id) – Reference to the product')
add_bullet(doc, 'product_name (VARCHAR) – Product name snapshot at time of order')
add_bullet(doc, 'product_price (DECIMAL) – Product price snapshot at time of order')
add_bullet(doc, 'quantity (INTEGER) – Quantity ordered')
add_bullet(doc, 'total_price (DECIMAL) – Line item total (price × quantity)')

add_sub_heading(doc, '8.3 Relationships')

add_body_text(doc, 'The following relationships exist between the entities in the AgroVeda database:')

rel_data = [
    ['Users → Cart_Items', '1:N (One-to-Many)', 'One user can have multiple items in their cart'],
    ['Users → Orders', '1:N (One-to-Many)', 'One user can place multiple orders'],
    ['Users → User_Roles', '1:N (One-to-Many)', 'One user can have multiple roles'],
    ['Products → Cart_Items', '1:N (One-to-Many)', 'One product can appear in multiple carts'],
    ['Products → Order_Items', '1:N (One-to-Many)', 'One product can appear in multiple orders'],
    ['Orders → Order_Items', '1:N (One-to-Many)', 'One order contains multiple order items'],
]

create_styled_table(doc, ['Relationship', 'Cardinality', 'Description'], rel_data, col_widths=[1.8, 1.5, 3.2])

add_formatted_paragraph(doc, '', font_size=12, space_after=6)

add_sub_heading(doc, '8.4 ER Diagram')

add_diagram_image(doc, 'diagrams/er_diagram.png', width_inches=6.0)
add_body_text(doc, 'Figure 8.1: Entity-Relationship Diagram for AgroVeda showing all six entities, their attributes, and the relationships between them.')

add_page_break(doc)


# ============================================================
# 17. CHAPTER 9: DATA FLOW DIAGRAM
# ============================================================
add_chapter_title(doc, '9', 'Data Flow Diagram')

add_sub_heading(doc, '9.1 Introduction')

add_body_text(doc, 'Data Flow Diagrams (DFDs) provide a graphical representation of the flow of data through the AgroVeda system. DFDs help in understanding how data moves between external entities, processes, and data stores within the application. The diagrams are presented at three levels of abstraction: Context Diagram (Level 0), Level 1 DFD (Major Processes), and Level 2 DFD (Process Decomposition).')

add_sub_heading(doc, '9.2 DFD Level 0 – Context Diagram')

add_body_text(doc, 'The Context Diagram (Level 0 DFD) provides the highest-level view of the AgroVeda system, showing it as a single process interacting with external entities:')

add_body_text(doc, 'External Entities:')
add_bold_bullet(doc, 'Customer: ', 'Provides registration data, login credentials, product search queries, cart updates, order details, and payment information. Receives product listings, cart contents, order confirmations, payment receipts, and order status updates.')
add_bold_bullet(doc, 'Admin: ', 'Provides admin credentials, product data (CRUD), order status updates, and user role assignments. Receives dashboard analytics, order reports, and user listings.')
add_bold_bullet(doc, 'Razorpay Payment Gateway: ', 'Receives payment order creation requests with amounts. Returns payment order IDs, payment status, and payment verification signatures.')

add_diagram_image(doc, 'diagrams/dfd_level0.png', width_inches=6.0)
add_body_text(doc, 'Figure 9.1: Context Diagram (DFD Level 0) for AgroVeda')

add_sub_heading(doc, '9.3 DFD Level 1 – Major Processes')

add_body_text(doc, 'The Level 1 DFD decomposes the AgroVeda system into its major processes:')

add_bold_bullet(doc, 'Process 1.0 – User Authentication: ', 'Handles user registration, login, JWT token generation, and role verification. Interacts with the Users and User_Roles data stores.')
add_bold_bullet(doc, 'Process 2.0 – Product Management: ', 'Manages the product catalog including browsing, searching, filtering, and CRUD operations (admin). Interacts with the Products data store.')
add_bold_bullet(doc, 'Process 3.0 – Cart Management: ', 'Handles shopping cart operations including adding items, updating quantities, removing items, and calculating totals. Interacts with Cart_Items and Products data stores.')
add_bold_bullet(doc, 'Process 4.0 – Order Processing: ', 'Manages order creation, payment processing (Razorpay integration), order status updates, and order history retrieval. Interacts with Orders, Order_Items, and Cart_Items data stores.')
add_bold_bullet(doc, 'Process 5.0 – Admin Dashboard: ', 'Provides administrative functions including product management, order monitoring, user management, and analytics. Interacts with all data stores.')

add_diagram_image(doc, 'diagrams/dfd_level1.png', width_inches=6.0)
add_body_text(doc, 'Figure 9.2: Level 1 Data Flow Diagram for AgroVeda')

add_sub_heading(doc, '9.4 DFD Level 2 – Order Processing Decomposition')

add_body_text(doc, 'The Level 2 DFD decomposes Process 4.0 (Order Processing) into sub-processes:')

add_bold_bullet(doc, 'Process 4.1 – Validate Order: ', 'Verifies cart contents, checks product availability, validates customer delivery information, and calculates final order total with any applicable discounts.')
add_bold_bullet(doc, 'Process 4.2 – Create Razorpay Order: ', 'Generates a Razorpay order with the calculated amount using the Razorpay API. Returns the order ID and key to the frontend for payment widget initialization.')
add_bold_bullet(doc, 'Process 4.3 – Verify Payment: ', 'Receives payment callback data from Razorpay, verifies the payment signature using HMAC SHA256, and updates the order payment status accordingly.')
add_bold_bullet(doc, 'Process 4.4 – Confirm Order: ', 'Creates the order record in the database, generates a unique order number, transfers cart items to order items, clears the cart, and sends order confirmation to the customer.')
add_bold_bullet(doc, 'Process 4.5 – Update Order Status: ', 'Allows admin to update order status through the lifecycle (pending → processing → shipped → delivered) and handles order cancellation requests.')

add_diagram_image(doc, 'diagrams/dfd_level2.png', width_inches=6.0)
add_body_text(doc, 'Figure 9.3: Level 2 DFD – Order Processing Decomposition')

add_page_break(doc)


# ============================================================
# 18. CHAPTER 10: DATABASE AND TABLE
# ============================================================
add_chapter_title(doc, '10', 'Database and Table')

add_sub_heading(doc, '10.1 Database Overview')

add_body_text(doc, 'The AgroVeda application uses PostgreSQL as its primary relational database management system in the production environment. PostgreSQL was chosen for its robustness, ACID compliance, support for complex queries, and excellent performance characteristics. The database schema consists of six tables that model the core entities of the e-commerce platform: users, user_roles, products, cart_items, orders, and order_items.')

add_body_text(doc, 'For development and testing purposes, an in-memory JSON storage system serves as a zero-configuration fallback, allowing developers to run the application without a PostgreSQL instance. The database schema includes proper indexing on frequently queried columns, foreign key constraints for referential integrity, and default values for common fields.')

add_sub_heading(doc, '10.2 Table: users')

add_body_text(doc, 'The users table stores authentication credentials for all registered users:')

users_table = [
    ['id', 'SERIAL', 'PRIMARY KEY', 'Auto-incrementing unique user identifier'],
    ['email', 'VARCHAR(255)', 'UNIQUE, NOT NULL', 'User email address for login'],
    ['password_hash', 'VARCHAR(255)', 'NOT NULL', 'bcrypt-hashed password (60 chars)'],
    ['created_at', 'TIMESTAMP', 'DEFAULT NOW()', 'Account registration timestamp'],
]

create_styled_table(doc, ['Field', 'Type', 'Constraints', 'Description'], users_table, col_widths=[1.5, 1.5, 1.5, 2.0])

add_formatted_paragraph(doc, '', font_size=6, space_after=6)

add_sub_heading(doc, '10.3 Table: user_roles')

add_body_text(doc, 'The user_roles table implements role-based access control by associating roles with users:')

roles_table = [
    ['id', 'SERIAL', 'PRIMARY KEY', 'Auto-incrementing unique identifier'],
    ['user_id', 'INTEGER', 'FK → users.id, NOT NULL', 'Reference to the user'],
    ['role', 'VARCHAR(50)', 'NOT NULL', 'Role name: customer or admin'],
]

create_styled_table(doc, ['Field', 'Type', 'Constraints', 'Description'], roles_table, col_widths=[1.5, 1.5, 1.5, 2.0])

add_formatted_paragraph(doc, '', font_size=6, space_after=6)

add_sub_heading(doc, '10.4 Table: products')

add_body_text(doc, 'The products table stores the complete agricultural product catalog with bilingual support:')

products_table = [
    ['id', 'SERIAL', 'PRIMARY KEY', 'Auto-incrementing product ID'],
    ['name', 'VARCHAR(255)', 'NOT NULL', 'Product name in English'],
    ['name_hindi', 'VARCHAR(255)', '', 'Product name in Hindi'],
    ['category', 'VARCHAR(100)', 'NOT NULL', 'Product category'],
    ['price', 'DECIMAL(10,2)', 'NOT NULL', 'Price in INR'],
    ['unit', 'VARCHAR(50)', '', 'Unit of measurement'],
    ['usage', 'TEXT', '', 'Usage instructions'],
    ['description', 'TEXT', '', 'Description in English'],
    ['description_hindi', 'TEXT', '', 'Description in Hindi'],
    ['rating', 'DECIMAL(2,1)', 'DEFAULT 0', 'Average rating (0-5)'],
    ['reviews', 'INTEGER', 'DEFAULT 0', 'Total review count'],
    ['image_url', 'VARCHAR(500)', '', 'Product image URL'],
    ['bulk_offers', 'TEXT', '', 'Bulk offer information'],
    ['created_at', 'TIMESTAMP', 'DEFAULT NOW()', 'Product creation date'],
]

create_styled_table(doc, ['Field', 'Type', 'Constraints', 'Description'], products_table, col_widths=[1.5, 1.5, 1.5, 2.0])

add_formatted_paragraph(doc, '', font_size=6, space_after=6)

add_sub_heading(doc, '10.5 Table: cart_items')

add_body_text(doc, 'The cart_items table manages the shopping cart for each authenticated user:')

cart_table = [
    ['id', 'SERIAL', 'PRIMARY KEY', 'Auto-incrementing identifier'],
    ['user_id', 'INTEGER', 'FK → users.id, NOT NULL', 'Reference to cart owner'],
    ['product_id', 'INTEGER', 'FK → products.id, NOT NULL', 'Reference to the product'],
    ['quantity', 'INTEGER', 'NOT NULL, DEFAULT 1', 'Product quantity in cart'],
]

create_styled_table(doc, ['Field', 'Type', 'Constraints', 'Description'], cart_table, col_widths=[1.5, 1.5, 1.5, 2.0])

add_formatted_paragraph(doc, '', font_size=6, space_after=6)

add_sub_heading(doc, '10.6 Table: orders')

add_body_text(doc, 'The orders table stores comprehensive order information including delivery and payment details:')

orders_table = [
    ['id', 'SERIAL', 'PRIMARY KEY', 'Auto-incrementing order ID'],
    ['user_id', 'INTEGER', 'FK → users.id', 'Reference to ordering user'],
    ['order_number', 'VARCHAR(50)', 'UNIQUE, NOT NULL', 'Human-readable order number'],
    ['customer_name', 'VARCHAR(255)', 'NOT NULL', 'Delivery recipient name'],
    ['customer_address', 'TEXT', 'NOT NULL', 'Delivery address'],
    ['customer_phone', 'VARCHAR(20)', 'NOT NULL', 'Contact phone number'],
    ['subtotal', 'DECIMAL(10,2)', 'NOT NULL', 'Subtotal before discount'],
    ['discount_amount', 'DECIMAL(10,2)', 'DEFAULT 0', 'Discount amount applied'],
    ['total_amount', 'DECIMAL(10,2)', 'NOT NULL', 'Final total after discount'],
    ['payment_method', 'VARCHAR(20)', 'NOT NULL', 'razorpay or cod'],
    ['status', 'VARCHAR(20)', 'DEFAULT pending', 'Order lifecycle status'],
    ['razorpay_order_id', 'VARCHAR(100)', '', 'Razorpay order reference'],
    ['razorpay_payment_id', 'VARCHAR(100)', '', 'Razorpay payment reference'],
    ['created_at', 'TIMESTAMP', 'DEFAULT NOW()', 'Order placement time'],
]

create_styled_table(doc, ['Field', 'Type', 'Constraints', 'Description'], orders_table, col_widths=[1.5, 1.5, 1.5, 2.0])

add_formatted_paragraph(doc, '', font_size=6, space_after=6)

add_sub_heading(doc, '10.7 Table: order_items')

add_body_text(doc, 'The order_items table stores individual line items for each order with price snapshots:')

order_items_table = [
    ['id', 'SERIAL', 'PRIMARY KEY', 'Auto-incrementing identifier'],
    ['order_id', 'INTEGER', 'FK → orders.id, NOT NULL', 'Reference to parent order'],
    ['product_id', 'INTEGER', 'FK → products.id', 'Reference to the product'],
    ['product_name', 'VARCHAR(255)', 'NOT NULL', 'Product name at order time'],
    ['product_price', 'DECIMAL(10,2)', 'NOT NULL', 'Unit price at order time'],
    ['quantity', 'INTEGER', 'NOT NULL', 'Quantity ordered'],
    ['total_price', 'DECIMAL(10,2)', 'NOT NULL', 'Line total (price × qty)'],
]

create_styled_table(doc, ['Field', 'Type', 'Constraints', 'Description'], order_items_table, col_widths=[1.5, 1.5, 1.5, 2.0])

add_page_break(doc)


# ============================================================
# 19. CHAPTER 11: I/O AND INTERFACE DESIGN
# ============================================================
add_chapter_title(doc, '11', 'I/O and Interface Design')

add_sub_heading(doc, '11.1 Interface Design Philosophy')

add_body_text(doc, 'The user interface of AgroVeda has been designed following modern web design principles with a focus on usability, accessibility, and responsiveness. The design philosophy adheres to the following core principles:')

add_bullet(doc, 'Mobile-First Design: All interfaces are designed for mobile screens first and progressively enhanced for larger screens using TailwindCSS responsive breakpoints.')
add_bullet(doc, 'Consistency: A unified design system using Shadcn/UI components ensures visual consistency across all pages with consistent color palette, typography, spacing, and interaction patterns.')
add_bullet(doc, 'Accessibility: Semantic HTML elements, proper ARIA labels, keyboard navigation support, and sufficient color contrast ratios ensure the application is accessible to users with disabilities.')
add_bullet(doc, 'Bilingual Support: All user-facing text elements support Hindi and English, with a prominent language toggle for easy switching.')
add_bullet(doc, 'Intuitive Navigation: A clear navigation hierarchy with breadcrumbs, consistent header/footer, and logical page flow guides users through the shopping experience effortlessly.')

add_sub_heading(doc, '11.2 Homepage')

add_body_text(doc, 'The homepage serves as the primary landing page and entry point for customers visiting AgroVeda. It features a prominent promotional banner slider showcasing seasonal offers, new arrivals, and featured product categories. Below the banner, a curated product grid displays popular and recommended agricultural products with images, prices in INR, star ratings, and quick "Add to Cart" buttons. The homepage also includes category navigation cards allowing users to jump directly to specific product categories.')

add_body_text(doc, 'Input: Page load request, language preference from localStorage')
add_body_text(doc, 'Output: Rendered homepage with banner, product grid, category cards, and navigation elements')

add_body_text(doc, '[SCREENSHOT PLACEHOLDER – Homepage]')

add_sub_heading(doc, '11.3 Product Catalog Page')

add_body_text(doc, 'The product catalog page displays the complete range of agricultural products available on AgroVeda. It features a sidebar filter panel with category checkboxes (Fertilizers, Pesticides, Growth Boosters, Herbicides, Organic), price range selection, and rating filter. The main content area displays products in a responsive grid layout that adapts from one column on mobile to four columns on desktop. Each product card shows the product image, name in the selected language, price, rating stars, review count, and an "Add to Cart" button.')

add_body_text(doc, 'Input: Category selection, price range, rating filter, search query, language toggle')
add_body_text(doc, 'Output: Filtered and sorted product grid with real-time updates')

add_body_text(doc, '[SCREENSHOT PLACEHOLDER – Product Catalog]')

add_sub_heading(doc, '11.4 Product Detail Page')

add_body_text(doc, 'The product detail page provides comprehensive information about a selected agricultural product. It displays a large product image, product name in both Hindi and English, detailed description, usage instructions, price with bulk offer information, customer rating and review count, and quantity selector with "Add to Cart" button. Related products from the same category are displayed at the bottom of the page.')

add_body_text(doc, 'Input: Product ID from URL parameters')
add_body_text(doc, 'Output: Complete product details with bilingual content, pricing, and related products')

add_body_text(doc, '[SCREENSHOT PLACEHOLDER – Product Detail]')

add_sub_heading(doc, '11.5 Cart Page')

add_body_text(doc, 'The cart page displays all items currently in the customer\'s shopping cart. Each cart item shows the product image, name, unit price, quantity controls (increment/decrement buttons), line total, and a remove button. The page includes a cart summary section showing subtotal, applicable discounts, delivery charges, and the final total amount. A prominent "Proceed to Checkout" button guides users to the next step.')

add_body_text(doc, 'Input: Cart item quantity changes, remove item actions, discount code application')
add_body_text(doc, 'Output: Updated cart with recalculated totals, navigation to checkout')

add_body_text(doc, '[SCREENSHOT PLACEHOLDER – Cart Page]')

add_sub_heading(doc, '11.6 Checkout Page')

add_body_text(doc, 'The checkout page collects delivery information and facilitates payment. It features a form with fields for customer name, delivery address (textarea), and phone number. The order summary section displays all cart items with quantities and prices. Payment method selection allows choosing between Razorpay (online payment) and Cash on Delivery. For Razorpay payments, clicking "Pay Now" opens the Razorpay checkout widget; for COD, clicking "Place Order" creates the order directly.')

add_body_text(doc, 'Input: Customer name, address, phone, payment method selection')
add_body_text(doc, 'Output: Order confirmation with order number, payment receipt (for Razorpay)')

add_body_text(doc, '[SCREENSHOT PLACEHOLDER – Checkout Page]')

add_sub_heading(doc, '11.7 Login / Signup Page')

add_body_text(doc, 'The authentication page provides a tabbed interface for Login and Sign Up functionality. The Login tab contains email and password fields with a "Login" button. The Sign Up tab contains email, password, and confirm password fields with a "Create Account" button. Form validation ensures email format correctness, minimum password length (6 characters), and password confirmation match. Error messages are displayed inline for invalid inputs.')

add_body_text(doc, 'Input: Email, password, confirm password (signup)')
add_body_text(doc, 'Output: JWT token stored in localStorage, redirect to homepage or admin dashboard')

add_body_text(doc, '[SCREENSHOT PLACEHOLDER – Login/Signup Page]')

add_sub_heading(doc, '11.8 Admin Dashboard')

add_body_text(doc, 'The admin dashboard provides an overview of platform operations with key metrics displayed in card format: Total Products, Total Orders, Total Revenue, and Pending Orders. A recent orders table shows the latest 10 orders with order number, customer name, total amount, status badge, and date. Navigation links in the sidebar provide access to Product Management, Order Management, and User Management sections.')

add_body_text(doc, 'Input: Admin authentication token')
add_body_text(doc, 'Output: Dashboard metrics, recent orders table, navigation to management sections')

add_body_text(doc, '[SCREENSHOT PLACEHOLDER – Admin Dashboard]')

add_sub_heading(doc, '11.9 Admin Product Management')

add_body_text(doc, 'The product management page displays all products in a sortable, searchable table with columns for ID, Name, Category, Price, Rating, and Actions (Edit/Delete). An "Add New Product" button opens a modal form with fields for all product attributes including bilingual fields (name, description in Hindi and English), category dropdown, price, unit, usage instructions, image URL, and bulk offers. The same form is used for editing existing products with pre-populated values.')

add_body_text(doc, 'Input: Product form data (name, category, price, descriptions, etc.)')
add_body_text(doc, 'Output: Created/updated product records, success/error notifications')

add_body_text(doc, '[SCREENSHOT PLACEHOLDER – Admin Product Management]')

add_page_break(doc)


# ============================================================
# 20. CHAPTER 12: TESTING
# ============================================================
add_chapter_title(doc, '12', 'Testing')

add_sub_heading(doc, '12.1 Testing Strategy')

add_body_text(doc, 'Testing is a critical phase in the software development life cycle that ensures the quality, reliability, and correctness of the AgroVeda application. A comprehensive testing strategy was employed that includes manual functional testing, integration testing, and user acceptance testing. The testing process was designed to validate all functional requirements, verify non-functional attributes (performance, security, usability), and ensure the application behaves correctly across different browsers and devices.')

add_body_text(doc, 'The testing approach followed a systematic methodology:')
add_bullet(doc, 'Unit-level verification of individual components and API endpoints')
add_bullet(doc, 'Integration testing of component interactions and API-database flows')
add_bullet(doc, 'System testing of end-to-end user workflows (browse → cart → checkout → payment)')
add_bullet(doc, 'Cross-browser compatibility testing on Chrome, Firefox, Edge, and Safari')
add_bullet(doc, 'Responsive design testing across mobile, tablet, and desktop viewports')
add_bullet(doc, 'Security testing of authentication, authorization, and input validation')

add_sub_heading(doc, '12.2 Test Environment')

add_body_text(doc, 'The following test environment was configured for executing the test cases:')

test_env = [
    ['Operating System', 'Windows 11 Pro'],
    ['Frontend Server', 'Vite Dev Server (port 5173)'],
    ['Backend API Server', 'Node.js Express (port 3001)'],
    ['Database', 'In-Memory JSON (dev) / PostgreSQL (staging)'],
    ['Primary Browser', 'Google Chrome 120+'],
    ['Secondary Browsers', 'Firefox 120+, Microsoft Edge 120+'],
    ['Mobile Testing', 'Chrome DevTools Device Emulation (iPhone 14, Samsung Galaxy S23)'],
    ['API Testing', 'Postman v10+'],
    ['Network', 'Local development network (localhost)'],
]

create_styled_table(doc, ['Component', 'Specification'], test_env, col_widths=[2.0, 4.5])

add_formatted_paragraph(doc, '', font_size=6, space_after=6)

add_sub_heading(doc, '12.3 Test Cases')

add_body_text(doc, 'The following test cases were executed to validate the functionality of the AgroVeda application:')

test_cases = [
    ['TC-01', 'Homepage Load', 'Navigate to homepage URL', 'Homepage loads with banner slider, product grid, and navigation. All images load correctly.', 'Pass'],
    ['TC-02', 'Product Filter', 'Select "Fertilizers" category from filter panel', 'Only fertilizer products are displayed. Other categories are hidden. Filter state persists.', 'Pass'],
    ['TC-03', 'Add to Cart', 'Click "Add to Cart" on a product while logged in', 'Product is added to cart. Cart badge count increments. Success notification shown.', 'Pass'],
    ['TC-04', 'Cart Quantity Update', 'Increase quantity of cart item from 1 to 3', 'Quantity updates to 3. Line total and cart total recalculate correctly.', 'Pass'],
    ['TC-05', 'Checkout with COD', 'Fill delivery form and select COD payment', 'Order is created with status "pending". Order confirmation page shown with order number.', 'Pass'],
    ['TC-06', 'Razorpay Payment', 'Select Razorpay payment and complete test payment', 'Razorpay widget opens. Test payment succeeds. Order status updates to "confirmed".', 'Pass'],
    ['TC-07', 'Admin Login', 'Login with admin credentials', 'Admin dashboard loads with metrics. Product/Order management links visible.', 'Pass'],
    ['TC-08', 'Admin CRUD', 'Create a new product via admin panel', 'Product is created and appears in product list. All fields saved correctly.', 'Pass'],
    ['TC-09', 'Order Status Update', 'Change order status from "pending" to "shipped"', 'Order status updates. Status badge color changes. Timestamp recorded.', 'Pass'],
    ['TC-10', 'Responsive Design', 'Test on iPhone 14 viewport (390×844)', 'Layout adapts to mobile. Navigation collapses to hamburger menu. Product grid shows single column.', 'Pass'],
]

# Create test case table
tc_table = doc.add_table(rows=1 + len(test_cases), cols=5)
tc_table.alignment = WD_TABLE_ALIGNMENT.CENTER

tc_headers = ['TC ID', 'Test Case', 'Input/Action', 'Expected Result', 'Status']
for i, h in enumerate(tc_headers):
    cell = tc_table.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(h)
    run.font.name = 'Arial'
    run.font.size = Pt(9)
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_shading(cell, "D9E2F3")

for r_idx, row_data in enumerate(test_cases):
    for c_idx, cell_text in enumerate(row_data):
        cell = tc_table.rows[r_idx + 1].cells[c_idx]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(cell_text))
        run.font.name = 'Arial'
        run.font.size = Pt(8)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx in [2, 3] else WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if c_idx == 4:
            if cell_text == 'Pass':
                set_cell_shading(cell, "C6EFCE")
            else:
                set_cell_shading(cell, "FFC7CE")

# Set column widths for test case table
tc_widths = [0.6, 1.0, 1.5, 2.2, 0.6]
for row in tc_table.rows:
    for i, width in enumerate(tc_widths):
        row.cells[i].width = Inches(width)

# Add borders to test case table
tbl = tc_table._tbl
tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>') 
borders = parse_xml(
    f'<w:tblBorders {nsdecls("w")}>'
    f'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
    f'<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
    f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
    f'<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
    f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
    f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
    f'</w:tblBorders>'
)
tblPr.append(borders)

add_formatted_paragraph(doc, '', font_size=6, space_after=6)

add_sub_heading(doc, '12.4 Test Results Summary')

add_body_text(doc, 'The following table summarizes the overall testing results for the AgroVeda application:')

results_summary = [
    ['Total Test Cases Executed', '10'],
    ['Test Cases Passed', '10'],
    ['Test Cases Failed', '0'],
    ['Pass Rate', '100%'],
    ['Defects Found During Testing', '3 (all resolved)'],
    ['Critical Defects', '0'],
    ['Regression Issues', '0'],
]

create_styled_table(doc, ['Metric', 'Value'], results_summary, col_widths=[3.0, 3.0])

add_formatted_paragraph(doc, '', font_size=6, space_after=6)

add_sub_heading(doc, '12.5 Defects Found and Resolved')

add_body_text(doc, 'During the testing phase, three minor defects were identified and resolved:')

add_bold_bullet(doc, 'DEF-001 (Low): ', 'Cart total displayed incorrect amount when discount code was applied with a percentage greater than 50%. Fixed by adding a maximum discount cap validation.')
add_bold_bullet(doc, 'DEF-002 (Medium): ', 'Hindi text rendering caused layout overflow on product cards in mobile view. Fixed by adding text truncation with ellipsis CSS properties and adjusting card minimum heights.')
add_bold_bullet(doc, 'DEF-003 (Low): ', 'Order status filter in admin panel did not reset when navigating away and returning. Fixed by resetting filter state on component unmount using useEffect cleanup.')

add_page_break(doc)


# ============================================================
# 21. CHAPTER 13: FUTURE SCOPE
# ============================================================
add_chapter_title(doc, '13', 'Future Scope')

add_sub_heading(doc, '13.1 Introduction')

add_body_text(doc, 'While the current version of AgroVeda successfully delivers a functional agricultural e-commerce platform with core shopping, payment, and administration features, there exist numerous opportunities for enhancement and expansion. The platform\'s modular architecture and modern technology stack provide a solid foundation for incorporating advanced features and scaling the application to serve a larger user base. The future development roadmap is organized into short-term, medium-term, and long-term goals.')

add_sub_heading(doc, '13.2 Short-Term Enhancements (3-6 Months)')

add_body_text(doc, 'The short-term development plan focuses on incremental improvements that enhance user experience and operational efficiency:')

add_bold_bullet(doc, 'Multi-Language Support: ', 'Expanding beyond Hindi and English to support regional languages such as Marathi, Gujarati, Tamil, and Telugu. This will involve implementing an i18n (internationalization) framework using libraries like react-i18next, enabling dynamic language switching and automatic content translation management.')

add_bold_bullet(doc, 'Real-Time Notifications via Twilio: ', 'Integrating Twilio\'s communication APIs for sending real-time SMS and WhatsApp notifications for order confirmations, shipping updates, delivery notifications, and promotional offers. This will significantly improve customer engagement and provide timely order status updates to farmers who may not have consistent email access.')

add_bold_bullet(doc, 'Enhanced Search with Elasticsearch: ', 'Implementing full-text search functionality using Elasticsearch to provide fast, fuzzy, and typo-tolerant product search across both English and Hindi product catalogs.')

add_bold_bullet(doc, 'Customer Reviews and Ratings System: ', 'Building a user-generated content system where customers can leave detailed reviews, upload photos, and rate products they have purchased, helping other buyers make informed decisions.')

add_sub_heading(doc, '13.3 Medium-Term Developments (6-12 Months)')

add_body_text(doc, 'The medium-term goals focus on adding intelligent features and scaling the business model:')

add_bold_bullet(doc, 'AI-Powered Crop Recommendation Engine: ', 'Developing a machine learning-based recommendation system that analyzes soil type, crop variety, season, region, and historical purchase data to suggest the most suitable fertilizers, pesticides, and growth boosters for individual farmers. This feature would leverage TensorFlow.js or integrate with a Python-based ML backend.')

add_bold_bullet(doc, 'Multi-Vendor Marketplace: ', 'Transforming AgroVeda from a single-vendor platform to a multi-vendor marketplace where multiple agricultural product manufacturers and distributors can register, list their products, set pricing, manage inventory, and process orders independently. This involves implementing vendor registration, commission management, vendor analytics, and dispute resolution workflows.')

add_bold_bullet(doc, 'Progressive Web App (PWA): ', 'Converting AgroVeda into a Progressive Web App with service worker-based offline caching, push notifications, and add-to-home-screen functionality, enabling farmers in areas with intermittent internet connectivity to browse products offline.')

add_sub_heading(doc, '13.4 Long-Term Vision (12-24 Months)')

add_body_text(doc, 'The long-term vision positions AgroVeda as a comprehensive agricultural technology platform:')

add_bold_bullet(doc, 'IoT Soil Sensor Integration: ', 'Partnering with IoT device manufacturers to integrate soil sensor data (pH level, moisture content, nitrogen/phosphorus/potassium levels) directly into the AgroVeda platform. This data would feed into the recommendation engine to provide highly personalized product suggestions based on real-time soil conditions.')

add_bold_bullet(doc, 'Drone Delivery Tracking: ', 'Implementing real-time delivery tracking with GPS integration and exploring drone-based delivery partnerships for remote agricultural areas where traditional logistics face challenges. The tracking system would provide live delivery status updates with estimated arrival times.')

add_bold_bullet(doc, 'Blockchain Supply Chain Transparency: ', 'Implementing a blockchain-based supply chain tracking system that provides complete transparency from product manufacturing to delivery. Customers would be able to verify product authenticity, manufacturing date, batch number, and supply chain journey through QR code scanning and blockchain verification, building trust and combating counterfeit agricultural products.')

add_page_break(doc)


# ============================================================
# 22. CHAPTER 14: CONCLUSION
# ============================================================
add_chapter_title(doc, '14', 'Conclusion')

add_sub_heading(doc, '14.1 Project Summary')

add_body_text(doc, 'The AgroVeda – E-Commerce Platform for Agricultural Products has been successfully designed, developed, and deployed as a comprehensive full-stack web application that addresses the challenges of agricultural product distribution in India. The project represents a significant effort in applying modern web development technologies to solve a real-world problem that affects millions of farmers and agricultural retailers across the country.')

add_body_text(doc, 'Over the course of approximately 20 working days organized into 6 Agile sprints, the project progressed from initial concept and database schema design to a fully functional, deployed application. The final product includes a responsive customer-facing storefront with bilingual (Hindi/English) support, a shopping cart and checkout system, Razorpay payment gateway integration, and a secure admin dashboard for product and order management.')

add_sub_heading(doc, '14.2 Objectives Achieved')

add_body_text(doc, 'All primary and secondary objectives defined at the project\'s inception have been successfully achieved:')

add_bullet(doc, 'A complete product catalog with bilingual support (Hindi and English) was implemented, allowing farmers to browse and understand product details in their preferred language.')
add_bullet(doc, 'A fully functional shopping cart system with quantity management, discount code support, and real-time total calculations was developed.')
add_bullet(doc, 'Razorpay payment gateway was successfully integrated, supporting UPI, debit/credit cards, and net banking, along with Cash on Delivery as a fallback option.')
add_bullet(doc, 'A secure admin dashboard with CRUD operations for product management, order tracking, and user role management was built with JWT-based role access control.')
add_bullet(doc, 'The application was deployed on Vercel with PostgreSQL database integration, achieving production-ready status with SSL encryption and global CDN distribution.')
add_bullet(doc, 'Responsive mobile-first design was implemented ensuring optimal user experience across smartphones, tablets, and desktop computers.')

add_sub_heading(doc, '14.3 Lessons Learned')

add_body_text(doc, 'The development of AgroVeda provided invaluable learning experiences across multiple dimensions:')

add_bold_bullet(doc, 'Full-Stack Architecture: ', 'Gained hands-on experience in designing and implementing a complete client-server architecture with clear separation of concerns between the React frontend and Express.js backend.')
add_bold_bullet(doc, 'TypeScript Benefits: ', 'Experienced firsthand how TypeScript\'s static type system catches errors during development, improves code documentation, and enhances IDE tooling for a more productive development experience.')
add_bold_bullet(doc, 'Payment Gateway Integration: ', 'Learned the intricacies of integrating third-party payment gateways including order creation flows, webhook handling, signature verification, and error recovery mechanisms.')
add_bold_bullet(doc, 'Database Design: ', 'Gained practical experience in relational database design, including normalization, foreign key constraints, indexing strategies, and the trade-offs between relational and in-memory storage systems.')
add_bold_bullet(doc, 'Deployment & DevOps: ', 'Learned about serverless deployment, environment variable management, build optimization, and the operational considerations of maintaining a production application.')

add_sub_heading(doc, '14.4 Documentation Importance')

add_body_text(doc, 'The process of creating this comprehensive project report has reinforced the importance of thorough documentation in software development. Technical documentation serves multiple critical purposes: it facilitates knowledge transfer, aids in onboarding future developers, provides a reference for system maintenance and debugging, and ensures that the rationale behind design decisions is preserved for posterity. The discipline of documenting each system component, from database schemas to user interface designs, has deepened the understanding of the system\'s architecture and highlighted areas for potential improvement.')

add_sub_heading(doc, '14.5 Personal Reflection')

add_body_text(doc, 'Working on the AgroVeda project has been a transformative experience that bridged the gap between academic learning and practical software development. The project provided opportunities to apply theoretical concepts from subjects such as Database Management Systems, Software Engineering, Web Technologies, and Computer Networks to a tangible, real-world application. The experience of working through the complete software development lifecycle – from requirements gathering and system design to implementation, testing, and deployment – has provided a holistic understanding of the software development profession.')

add_sub_heading(doc, '14.6 Closing Note')

add_body_text(doc, 'AgroVeda stands as a testament to the potential of technology to transform traditional industries and improve the lives of people in rural communities. By providing a digital platform for agricultural product distribution, the project contributes to the broader vision of Digital India and demonstrates how modern web technologies can be leveraged to address real-world challenges in the agricultural sector. The platform\'s modular architecture and scalable design ensure that it can evolve and grow to meet the changing needs of the agricultural community, ultimately contributing to improved agricultural productivity and farmer welfare in India.')

add_page_break(doc)


# ============================================================
# 23. CHAPTER 15: REFERENCES
# ============================================================
add_chapter_title(doc, '15', 'References')

add_body_text(doc, 'The following resources were referenced during the development of the AgroVeda project and the preparation of this report:')

add_formatted_paragraph(doc, '', font_size=6, space_after=6)

references = [
    ['1', 'React Documentation', 'https://react.dev/', 'Official documentation for React 18, covering components, hooks, state management, and best practices.'],
    ['2', 'Node.js Documentation', 'https://nodejs.org/docs/', 'Official documentation for Node.js runtime environment, including API references and guides.'],
    ['3', 'Vite Documentation', 'https://vitejs.dev/guide/', 'Official documentation for Vite build tool, covering configuration, plugins, and deployment.'],
    ['4', 'Razorpay Integration Guide', 'https://razorpay.com/docs/', 'Comprehensive documentation for Razorpay payment gateway integration including API references, webhooks, and SDKs.'],
    ['5', 'PostgreSQL Documentation', 'https://www.postgresql.org/docs/', 'Official documentation for PostgreSQL database system, including SQL reference, data types, and administration guides.'],
    ['6', 'Vercel Documentation', 'https://vercel.com/docs', 'Official documentation for Vercel deployment platform, covering frameworks, serverless functions, and environment configuration.'],
    ['7', 'TailwindCSS Documentation', 'https://tailwindcss.com/docs', 'Official documentation for TailwindCSS utility-first CSS framework, including responsive design and customization.'],
    ['8', 'JWT.io', 'https://jwt.io/', 'JSON Web Token specification, debugger, and library references for implementing JWT-based authentication.'],
    ['9', 'MDN Web Docs', 'https://developer.mozilla.org/', 'Mozilla Developer Network comprehensive web development documentation including HTML, CSS, JavaScript, and Web APIs.'],
    ['10', 'TypeScript Handbook', 'https://www.typescriptlang.org/docs/', 'Official TypeScript documentation covering type system, configuration, and migration guides.'],
    ['11', 'Express.js Guide', 'https://expressjs.com/en/guide/', 'Official Express.js web framework documentation including routing, middleware, and error handling.'],
    ['12', 'Shadcn/UI Documentation', 'https://ui.shadcn.com/', 'Documentation for Shadcn/UI component library built on Radix UI primitives.'],
]

ref_table = doc.add_table(rows=1 + len(references), cols=4)
ref_table.alignment = WD_TABLE_ALIGNMENT.CENTER

ref_headers = ['S.No.', 'Reference', 'URL', 'Description']
for i, h in enumerate(ref_headers):
    cell = ref_table.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(h)
    run.font.name = 'Arial'
    run.font.size = Pt(9)
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_shading(cell, "D9E2F3")

for r_idx, row_data in enumerate(references):
    for c_idx, cell_text in enumerate(row_data):
        cell = ref_table.rows[r_idx + 1].cells[c_idx]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(str(cell_text))
        run.font.name = 'Arial'
        run.font.size = Pt(8)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx in [2, 3] else WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if c_idx == 2:
            run.font.color.rgb = RGBColor(0, 0, 238)
            run.font.underline = True

# Set column widths for references table
ref_widths = [0.5, 1.3, 2.2, 2.5]
for row in ref_table.rows:
    for i, width in enumerate(ref_widths):
        row.cells[i].width = Inches(width)

# Add borders
tbl = ref_table._tbl
tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>') 
borders = parse_xml(
    f'<w:tblBorders {nsdecls("w")}>'
    f'<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
    f'<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
    f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
    f'<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
    f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
    f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
    f'</w:tblBorders>'
)
tblPr.append(borders)


# ============================================================
# REMOVE PAGE NUMBERS FROM FIRST SECTION (Cover + Certificates)
# ============================================================
# The first section should have no page numbers
first_section = doc.sections[0]
footer1 = first_section.footer
# Clear footer in first section
for p in footer1.paragraphs:
    p.clear()

# Also make sure first section doesn't link to second section's footer
footer1.is_linked_to_previous = False


# ============================================================
# SAVE DOCUMENT
# ============================================================
output_path = r'd:\Study\Projects\Minor Project\New folder\Zoya_Ibrahim_Major_Project_Report_Final.docx'
doc.save(output_path)
print(f"Document saved successfully: {output_path}")
print(f"Total sections: {len(doc.sections)}")
print(f"Document generated with all 15 chapters + certificates + front matter")
